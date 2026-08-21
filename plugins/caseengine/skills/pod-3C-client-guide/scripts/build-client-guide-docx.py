#!/usr/bin/env python3
"""
What: Builds a CE-branded Client Guide DOCX from a client-guide-data.json file, plus a
      paired .md sibling at the same base path (both are deliverables).
Input: client-guide-data.json (shape per the canonical schema:
       references/schemas/client-guide.json)
       + a CE logo image (download from the canonical Case Engine Branding folder:
       https://drive.google.com/drive/folders/1OulNcg6hZ6caFD0Xc-W50i7ougC_y0qo)
Output: Client Guide.docx (CE-branded, cover page + headers/footers + Episode Overview +
        Pre-Interview Prep + Segment Breakdown + FAQ) AND Client Guide.md (raw markdown
        deliverable, sibling at the same base path).
Re-run: Safe - overwrites both outputs. Drive's auto-conversion preserves the same fileId
        when the calling skill uses files.update for re-uploads (per the Push to Drive doc).

Convention sync: section order, branding spec, and the canonical filename live in the
Podcast Drive doc + the Case Engine Branding folder. If those change, update this script
to match. The canonical sources are the single source of truth.

Hard rules:
- NEVER emit an "Internal Setup" / "Complete and delete this section before sharing"
  checklist. The new pipeline configures the deliverable correctly from the jump.
- STRIP pandoc artifacts in input data: `[text]{.underline}` becomes a DOCX underline run
  (in DOCX output) or just `text` (in MD output). Same treatment for `{.smallcaps}`,
  `{.mark}`, `{.color=...}` patterns.
- Section order is fixed: Episode Overview -> Pre-Interview Prep -> Segment Breakdown -> FAQ.
- Cover order is fixed (2026-06-17 fix): "Client Guide" title -> EPISODE TITLE as the
  prominent subtitle (MANDATORY) -> FIRM NAME as a secondary line -> "{practice area} |
  {scope - location}" -> "Prepared by Case Engine" + date. The episode title is the
  subtitle, NEVER the firm name (this corrects the prior cover that emitted the firm as
  the subtitle and omitted the episode title - the Sutliff E8 defect).

Dependencies: python-docx (cowork runtime has it). No optional deps.

Usage:
  python3 build-client-guide-docx.py \\
      --data /path/to/client-guide-data.json \\
      --logo /path/to/ce-logo.png \\
      --firm "The May Firm" \\
      --attorney "Robert May" \\
      --practice-area "Car Accidents" \\
      --episode-title "How to File a Car Accident Claim" \\
      --scope "Location" \\
      --output "/path/to/Client Guide - E2 - How to File a Car Accident Claim - CA.docx" \\
      [--location "California"]            (omit at Topic Only scope)
      [--run-date "April 27, 2026"]        (defaults to today)
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# CE brand (per Case Engine Branding folder colors.md)
CE_BLUE = RGBColor(0x35, 0x73, 0xFF)
CE_DARK = RGBColor(0x0F, 0x17, 0x2A)
CE_GRAY = RGBColor(0x5B, 0x66, 0x76)
FONT = "Roboto"  # CE branded deliverable font (overrides Calibri per Gabe 2026-05-12; if the Branding folder spec still says Calibri, Roboto wins)


# ---------- Pandoc artifact stripping ----------

# Matches Pandoc bracketed-span markup like [text]{.underline}, [text]{.smallcaps},
# [text]{.mark}, [text]{.color=red}, [text]{.attr1 .attr2 key=value}, etc.
_PANDOC_SPAN_RE = re.compile(r"\[([^\[\]]*?)\]\{[^{}]*?\}")


def parse_pandoc_runs(text):
    """Split a string into a list of (text, attrs_dict) pairs.

    `attrs_dict` is a dict of inline-style flags parsed from the pandoc bracketed-span
    markup ({.underline}, {.smallcaps}, {.mark}, {.color=...}). Plain text returns
    a single tuple with empty attrs.

    Examples:
      "Plain"                         -> [("Plain", {})]
      "Hi [Bob]{.underline}"          -> [("Hi ", {}), ("Bob", {"underline": True})]
      "[A]{.smallcaps} and [B]{.mark}" -> [("A", {"smallcaps": True}), (" and ", {}),
                                           ("B", {"mark": True})]

    The build script uses this to render proper underline/highlight runs in DOCX, and
    to strip-but-keep the inner text in MD output.
    """
    if text is None:
        return [("", {})]
    parts = []
    cursor = 0
    for m in _PANDOC_SPAN_RE.finditer(text):
        if m.start() > cursor:
            parts.append((text[cursor:m.start()], {}))
        inner = m.group(1)
        attr_str = m.group(0)[m.end(1) - m.start() + 2:-1]  # contents between { and }
        attrs = _parse_pandoc_attrs(attr_str)
        parts.append((inner, attrs))
        cursor = m.end()
    if cursor < len(text):
        parts.append((text[cursor:], {}))
    if not parts:
        parts.append((text, {}))
    return parts


def _parse_pandoc_attrs(attr_str):
    """Parse a pandoc attribute string body into a flag dict.

    Recognizes: .underline, .smallcaps, .mark, .underline, .color=NAME, color=NAME.
    Anything unrecognized is ignored (no-op). Multiple attrs separated by spaces.
    """
    attrs = {}
    tokens = attr_str.strip().split()
    for tok in tokens:
        tok = tok.strip()
        if tok in (".underline", "underline"):
            attrs["underline"] = True
        elif tok in (".smallcaps", "smallcaps"):
            attrs["smallcaps"] = True
        elif tok in (".mark", "mark"):
            attrs["mark"] = True
        elif tok.startswith(".color=") or tok.startswith("color="):
            attrs["color"] = tok.split("=", 1)[1]
    return attrs


def strip_pandoc(text):
    """Strip pandoc bracketed-span wrappers and keep only the inner text.

    Used for the .md output (markdown can't represent pandoc spans natively).
    """
    if text is None:
        return ""
    return _PANDOC_SPAN_RE.sub(r"\1", text)


# ---------- DOCX style helpers ----------

def set_run_style(run, *, font=FONT, size=11, bold=False, italic=False, underline=False,
                  color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)


def add_para(doc, text="", *, align=None, space_before=0, space_after=6, **run_kw):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_style(run, **run_kw)
    return p


def add_pandoc_aware_para(doc, text, *, style=None, base_kw=None):
    """Add a paragraph that respects pandoc bracketed-span markup in the input.

    Renders [text]{.underline} as a DOCX underline run, {.smallcaps} as small-caps
    (approximated via uppercased run since python-docx has no smallcaps toggle),
    {.mark} as highlighted (yellow background approximated by color), and any
    `.color=NAME` as a colored run.
    """
    base_kw = base_kw or {}
    if style is not None:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    for chunk_text, attrs in parse_pandoc_runs(text):
        if not chunk_text:
            continue
        kw = dict(base_kw)
        if attrs.get("underline"):
            kw["underline"] = True
        if attrs.get("smallcaps"):
            chunk_text = chunk_text.upper()
        if attrs.get("mark"):
            kw["color"] = CE_BLUE
        run = p.add_run(chunk_text)
        set_run_style(run, **kw)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_header(section, text):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_style(run, size=9, italic=True, color=CE_GRAY)


def set_footer(section, text_left):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text_left + "\t\t")
    set_run_style(run, size=9, color=CE_GRAY)
    run2 = p.add_run("Page ")
    set_run_style(run2, size=9, color=CE_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_style(run, size=16, bold=True, color=CE_BLUE)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_style(run, size=13, bold=True, color=CE_DARK)
    return p


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_style(run, size=11, bold=True, color=CE_DARK)
    return p


def add_bullet_lead_detail(doc, lead, detail):
    """Render a bulleted item with bolded lead sentence + regular-weight detail.

    Pandoc artifacts inside `lead` and `detail` are stripped/applied as inline runs.
    """
    p = doc.add_paragraph(style="List Bullet")
    # Lead - bold
    for chunk_text, attrs in parse_pandoc_runs(lead):
        if not chunk_text:
            continue
        kw = {"size": 11, "bold": True, "color": CE_DARK}
        if attrs.get("underline"):
            kw["underline"] = True
        if attrs.get("smallcaps"):
            chunk_text = chunk_text.upper()
        run = p.add_run(chunk_text)
        set_run_style(run, **kw)
    # Space between lead and detail
    if detail:
        space = p.add_run(" ")
        set_run_style(space, size=11, color=CE_DARK)
        for chunk_text, attrs in parse_pandoc_runs(detail):
            if not chunk_text:
                continue
            kw = {"size": 11, "color": CE_DARK}
            if attrs.get("underline"):
                kw["underline"] = True
            if attrs.get("smallcaps"):
                chunk_text = chunk_text.upper()
            run = p.add_run(chunk_text)
            set_run_style(run, **kw)
    return p


def add_episode_plan_outline(doc, plan_items):
    """Render the episode_plan as an indentation-aware bullet list.

    Items with leading whitespace become indented (nested) bullets. Pandoc artifacts
    are stripped/applied. Plain bulleting via List Bullet style; nesting controlled
    by paragraph_format.left_indent.
    """
    for item in plan_items:
        # Detect leading whitespace as indent depth (2 spaces = 1 level).
        stripped = item.lstrip()
        indent_chars = len(item) - len(stripped)
        depth = max(0, indent_chars // 2)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 + 0.25 * depth)
        for chunk_text, attrs in parse_pandoc_runs(stripped):
            if not chunk_text:
                continue
            kw = {"size": 11, "color": CE_DARK}
            if attrs.get("underline"):
                kw["underline"] = True
            if attrs.get("smallcaps"):
                chunk_text = chunk_text.upper()
            run = p.add_run(chunk_text)
            set_run_style(run, **kw)


# ---------- Build ----------

def build(args):
    data_path = Path(args.data)
    if not data_path.exists():
        print(f"ERROR: {data_path} not found", file=sys.stderr)
        sys.exit(1)
    data = json.loads(data_path.read_text())

    firm = args.firm
    attorney = args.attorney
    practice_area = args.practice_area
    scope = args.scope
    location = args.location
    run_date = args.run_date or date.today().strftime("%B %d, %Y").replace(" 0", " ")
    host = data.get("host", "your co-host")

    overview = data.get("episode_overview", {})
    metadata = overview.get("metadata", {})
    episode_topic = data.get("episode_topic") or metadata.get("episode_topic", "")
    # Episode title for the cover subtitle (mandatory). Prefer the explicit --episode-title
    # arg; fall back to the data payload's episode_title, then episode_topic. Per the
    # 2026-06-17 cover fix, this is the prominent cover subtitle - never the firm name.
    episode_title = (args.episode_title
                     or data.get("episode_title")
                     or metadata.get("episode_title")
                     or episode_topic)
    if not episode_title:
        print("ERROR: episode title is mandatory on the cover - pass --episode-title "
              "or include episode_title/episode_topic in the data JSON", file=sys.stderr)
        sys.exit(1)
    duration_str = metadata.get("estimated_duration", "")
    question_count = metadata.get("question_count", 0)
    segment_count = metadata.get("segment_count", 0)
    episode_plan = overview.get("episode_plan", [])
    value_prop = overview.get("value_prop", "")

    prep = data.get("pre_interview_prep", {})
    things_to_think = prep.get("things_to_think_about", [])
    things_to_do = prep.get("things_to_do", [])

    segments = data.get("segments", {})
    intro_desc = segments.get("intro_description", "")
    intro_dur = segments.get("intro_duration", "~2 minutes")
    seg_list = segments.get("list", [])
    outro_desc = segments.get("outro_description", "")
    outro_dur = segments.get("outro_duration", "~2 minutes")

    faq = data.get("faq", [])

    scope_label = scope + (f" - {location}" if location else "")
    header_text = f"Case Engine  |  Client Guide  |  {firm}, {practice_area}"

    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        set_header(section, header_text)
        set_footer(section, "Case Engine  |  Confidential")

    # ---- Cover page ----
    for _ in range(3):
        doc.add_paragraph()

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(args.logo, width=Inches(2.4))

    # Cover order (per 2026-06-17 fix): title -> EPISODE TITLE (prominent subtitle,
    # mandatory) -> FIRM NAME (secondary line, kept) -> practice area | location ->
    # "Prepared by Case Engine" + date. Previously this emitted the firm as the subtitle
    # and omitted the episode title entirely (the Sutliff E8 cover defect). The episode
    # title is the prominent subtitle now; the firm is a secondary line below it.
    add_para(doc, "Client Guide", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=6, size=36, bold=True, color=CE_BLUE)
    # Prominent subtitle = EPISODE TITLE (mandatory)
    add_para(doc, episode_title, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=22, bold=True, color=CE_DARK)
    # Secondary line = CLIENT / FIRM NAME (kept on the cover)
    add_para(doc, firm, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=16, bold=True, color=CE_DARK)
    add_para(doc, f"{practice_area}  |  {scope_label}",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24, size=14, color=CE_DARK)
    add_para(doc, "Prepared by Case Engine", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=2, size=11, italic=True, color=CE_GRAY)
    add_para(doc, run_date, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=0, size=11, color=CE_GRAY)

    add_page_break(doc)

    # ---- Episode Overview ----
    add_h2(doc, "Episode Overview")
    if value_prop:
        add_pandoc_aware_para(doc, value_prop, base_kw={"size": 11, "color": CE_DARK})

    add_h3(doc, "Metadata")
    # Render as bold key + value lines
    meta_lines = []
    if episode_topic:
        meta_lines.append(("Episode Topic:", episode_topic))
    if duration_str:
        if question_count and segment_count:
            meta_lines.append(("Estimated Duration:",
                               f"{duration_str} ({question_count} questions across {segment_count} segments)"))
        else:
            meta_lines.append(("Estimated Duration:", duration_str))
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label + " ")
        set_run_style(run, size=11, bold=True, color=CE_DARK)
        for chunk_text, attrs in parse_pandoc_runs(value):
            if not chunk_text:
                continue
            kw = {"size": 11, "color": CE_DARK}
            if attrs.get("underline"):
                kw["underline"] = True
            run = p.add_run(chunk_text)
            set_run_style(run, **kw)
    add_para(doc,
             "Duration is always a range. Episodes run shorter or longer depending on how the conversation flows, and some topics get extension episodes.",
             size=11, italic=True, color=CE_GRAY, space_before=4, space_after=8)

    add_h3(doc, "Episode Plan")
    add_episode_plan_outline(doc, episode_plan)

    add_page_break(doc)

    # ---- Pre-Interview Prep ----
    add_h2(doc, "Pre-Interview Prep")

    add_h3(doc, "Things to Think About")
    for item in things_to_think:
        add_bullet_lead_detail(doc, item.get("lead", ""), item.get("detail", ""))

    add_h3(doc, "Things to Do")
    for item in things_to_do:
        add_bullet_lead_detail(doc, item.get("lead", ""), item.get("detail", ""))

    add_page_break(doc)

    # ---- Segment Breakdown ----
    add_h2(doc, "Segment Breakdown")
    add_para(doc,
             "Below is a breakdown of each segment with the topics you should be prepared to discuss. You don't need to memorize answers - just be familiar with the topics so the conversation flows naturally.",
             size=11, italic=True, color=CE_GRAY, space_after=10)

    # Intro
    add_h3(doc, f"Intro ({intro_dur})")
    add_pandoc_aware_para(doc, intro_desc,
                          base_kw={"size": 11, "italic": True, "color": CE_GRAY})

    # Continuous question numbering across segments
    q_counter = 1
    for idx, seg in enumerate(seg_list, start=1):
        seg_name = seg.get("name", f"Segment {idx}")
        seg_desc = seg.get("description", "")
        seg_dur = seg.get("duration", "")
        seg_questions = seg.get("questions", [])

        add_h3(doc, f"S{idx}: {seg_name} ({seg_dur})")
        add_pandoc_aware_para(doc, seg_desc,
                              base_kw={"size": 11, "italic": True, "color": CE_GRAY})

        if seg_questions:
            add_h4(doc, "Questions")
            for q in seg_questions:
                # Schema v1.1: q is {q_text, attorney_response_bullets: [{label, detail}]}
                # Schema v1.0 (legacy): q is a plain string. Detect and adapt.
                if isinstance(q, dict):
                    q_text = q.get("q_text", "")
                    bullets = q.get("attorney_response_bullets", []) or []
                else:
                    q_text = q
                    bullets = []
                # Parent Q bullet (List Bullet, indent level 1)
                p = doc.add_paragraph(style="List Bullet")
                run_q = p.add_run(f"Q{q_counter}: ")
                set_run_style(run_q, size=11, bold=True, color=CE_DARK)
                for chunk_text, attrs in parse_pandoc_runs(q_text):
                    if not chunk_text:
                        continue
                    kw = {"size": 11, "color": CE_DARK}
                    if attrs.get("underline"):
                        kw["underline"] = True
                    if attrs.get("smallcaps"):
                        chunk_text = chunk_text.upper()
                    run = p.add_run(chunk_text)
                    set_run_style(run, **kw)
                # Attorney-response sub-bullets (List Bullet 2 = indented one level).
                # Each sub-bullet renders as: `**Label:** detail` with pandoc runs preserved.
                for sb in bullets:
                    label = sb.get("label", "") if isinstance(sb, dict) else ""
                    detail = sb.get("detail", "") if isinstance(sb, dict) else ""
                    if not label and not detail:
                        continue
                    sp = doc.add_paragraph(style="List Bullet 2")
                    # Render label as bold (pandoc runs respected)
                    for chunk_text, attrs in parse_pandoc_runs(label):
                        if not chunk_text:
                            continue
                        kw = {"size": 11, "bold": True, "color": CE_DARK}
                        if attrs.get("underline"):
                            kw["underline"] = True
                        if attrs.get("smallcaps"):
                            chunk_text = chunk_text.upper()
                        run = sp.add_run(chunk_text)
                        set_run_style(run, **kw)
                    # Space between label and detail
                    if detail:
                        if label:
                            space = sp.add_run(" ")
                            set_run_style(space, size=11, color=CE_DARK)
                        # Render detail as regular weight (pandoc runs respected)
                        for chunk_text, attrs in parse_pandoc_runs(detail):
                            if not chunk_text:
                                continue
                            kw = {"size": 11, "color": CE_DARK}
                            if attrs.get("underline"):
                                kw["underline"] = True
                            if attrs.get("smallcaps"):
                                chunk_text = chunk_text.upper()
                            run = sp.add_run(chunk_text)
                            set_run_style(run, **kw)
                q_counter += 1

    # Outro
    add_h3(doc, f"Outro ({outro_dur})")
    add_pandoc_aware_para(doc, outro_desc,
                          base_kw={"size": 11, "italic": True, "color": CE_GRAY})

    add_page_break(doc)

    # ---- FAQ ----
    add_h2(doc, "FAQ")
    for item in faq:
        question = item.get("question", "")
        answer = item.get("answer", "")
        p = doc.add_paragraph(style="List Bullet")
        # Bold question
        for chunk_text, attrs in parse_pandoc_runs(question):
            if not chunk_text:
                continue
            kw = {"size": 11, "bold": True, "color": CE_DARK}
            if attrs.get("underline"):
                kw["underline"] = True
            run = p.add_run(chunk_text)
            set_run_style(run, **kw)
        # Space + answer
        if answer:
            space = p.add_run(" ")
            set_run_style(space, size=11, color=CE_DARK)
            for chunk_text, attrs in parse_pandoc_runs(answer):
                if not chunk_text:
                    continue
                kw = {"size": 11, "color": CE_DARK}
                if attrs.get("underline"):
                    kw["underline"] = True
                run = p.add_run(chunk_text)
                set_run_style(run, **kw)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved DOCX: {output_path}")

    # Also emit a sibling .md alongside the .docx with the same content shape.
    # The .md is the raw markdown deliverable Drive will upload as `text/markdown` (no auto-conversion).
    # The .docx is the Drive-rendered Google Doc sibling. Both live side-by-side in the scope folder.
    md_path = output_path.with_suffix(".md")
    md_text = build_markdown(
        firm=firm, attorney=attorney, practice_area=practice_area, scope=scope,
        location=location, run_date=run_date, host=host,
        episode_topic=episode_topic, duration_str=duration_str,
        question_count=question_count, segment_count=segment_count,
        value_prop=value_prop, episode_plan=episode_plan,
        things_to_think=things_to_think, things_to_do=things_to_do,
        intro_desc=intro_desc, intro_dur=intro_dur,
        seg_list=seg_list, outro_desc=outro_desc, outro_dur=outro_dur,
        faq=faq,
    )
    md_path.write_text(md_text)
    print(f"Saved MD:   {md_path}")
    print(f"Data: {len(things_to_think)} think-prompts, {len(things_to_do)} do-tasks, "
          f"{len(seg_list)} segments, {sum(len(s.get('questions', [])) for s in seg_list)} questions, "
          f"{len(faq)} FAQs")


def build_markdown(*, firm, attorney, practice_area, scope, location, run_date, host,
                   episode_topic, duration_str, question_count, segment_count,
                   value_prop, episode_plan,
                   things_to_think, things_to_do,
                   intro_desc, intro_dur, seg_list, outro_desc, outro_dur, faq):
    """Build the raw .md deliverable - same content shape as the DOCX, plain markdown.

    The .md ships alongside the .docx in Drive (uploaded as text/markdown without conversion).
    Downstream skills that prefer raw markdown can read the .md directly; humans can open
    the Google Doc sibling.

    Pandoc artifacts in the source data are stripped (markdown can't represent
    [text]{.underline} natively); the inner text is preserved.
    """
    lines = []

    # ---- Episode Overview ----
    lines.append("## Episode Overview")
    lines.append("")
    if value_prop:
        lines.append(strip_pandoc(value_prop))
        lines.append("")

    lines.append("### Metadata")
    lines.append("")
    if episode_topic:
        lines.append(f"**Episode Topic:** {strip_pandoc(episode_topic)}")
        lines.append("")  # Blank line forces a paragraph break — without it, Drive's MD→Doc converter collapses adjacent bold-label lines into one paragraph.
    if duration_str:
        if question_count and segment_count:
            lines.append(f"**Estimated Duration:** {duration_str} "
                         f"({question_count} questions across {segment_count} segments)")
        else:
            lines.append(f"**Estimated Duration:** {duration_str}")
    lines.append("")
    lines.append("Duration is always a range. Episodes run shorter or longer depending on how the conversation flows, and some topics get extension episodes.")
    lines.append("")

    lines.append("### Episode Plan")
    lines.append("")
    for item in episode_plan:
        stripped = item.lstrip()
        indent_chars = len(item) - len(stripped)
        depth = max(0, indent_chars // 2)
        prefix = "  " * depth + "- "
        lines.append(prefix + strip_pandoc(stripped))
    lines.append("")

    lines.append("---")
    lines.append("")

    # ---- Pre-Interview Prep ----
    lines.append("## Pre-Interview Prep")
    lines.append("")

    lines.append("### Things to Think About")
    lines.append("")
    for item in things_to_think:
        lead = strip_pandoc(item.get("lead", ""))
        detail = strip_pandoc(item.get("detail", ""))
        if detail:
            lines.append(f"- **{lead}** {detail}")
        else:
            lines.append(f"- **{lead}**")
    lines.append("")

    lines.append("### Things to Do")
    lines.append("")
    for item in things_to_do:
        lead = strip_pandoc(item.get("lead", ""))
        detail = strip_pandoc(item.get("detail", ""))
        if detail:
            lines.append(f"- **{lead}** {detail}")
        else:
            lines.append(f"- **{lead}**")
    lines.append("")

    lines.append("---")
    lines.append("")

    # ---- Segment Breakdown ----
    lines.append("## Segment Breakdown")
    lines.append("")
    lines.append("_Below is a breakdown of each segment with the topics you should be prepared to discuss. You don't need to memorize answers - just be familiar with the topics so the conversation flows naturally._")
    lines.append("")

    lines.append(f"### Intro ({intro_dur})")
    lines.append("")
    if intro_desc:
        lines.append(f"_{strip_pandoc(intro_desc)}_")
        lines.append("")

    q_counter = 1
    for idx, seg in enumerate(seg_list, start=1):
        seg_name = seg.get("name", f"Segment {idx}")
        seg_desc = seg.get("description", "")
        seg_dur = seg.get("duration", "")
        seg_questions = seg.get("questions", [])

        lines.append(f"### S{idx}: {seg_name} ({seg_dur})")
        lines.append("")
        if seg_desc:
            lines.append(f"_{strip_pandoc(seg_desc)}_")
            lines.append("")
        if seg_questions:
            lines.append("#### Questions")
            lines.append("")
            for q in seg_questions:
                # v1.1 dict shape vs v1.0 legacy string
                if isinstance(q, dict):
                    q_text = q.get("q_text", "")
                    bullets = q.get("attorney_response_bullets", []) or []
                else:
                    q_text = q
                    bullets = []
                lines.append(f"- Q{q_counter}: {strip_pandoc(q_text)}")
                # Indented sub-bullets carrying the attorney coverage checklist verbatim
                # from the Client ROS. Markdown can't render the pandoc underline runs, so
                # strip them to plain text - the DOCX sibling preserves the formatting.
                for sb in bullets:
                    if not isinstance(sb, dict):
                        continue
                    label = strip_pandoc(sb.get("label", ""))
                    detail = strip_pandoc(sb.get("detail", ""))
                    if label and detail:
                        lines.append(f"  - **{label}** {detail}")
                    elif label:
                        lines.append(f"  - **{label}**")
                    elif detail:
                        lines.append(f"  - {detail}")
                q_counter += 1
            lines.append("")

    lines.append(f"### Outro ({outro_dur})")
    lines.append("")
    if outro_desc:
        lines.append(f"_{strip_pandoc(outro_desc)}_")
        lines.append("")

    lines.append("---")
    lines.append("")

    # ---- FAQ ----
    lines.append("## FAQ")
    lines.append("")
    for item in faq:
        q = strip_pandoc(item.get("question", ""))
        a = strip_pandoc(item.get("answer", ""))
        if a:
            lines.append(f"- **{q}** {a}")
        else:
            lines.append(f"- **{q}**")
    lines.append("")

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Build CE-branded Client Guide DOCX (and paired .md) from client-guide-data.json")
    parser.add_argument("--data", required=True,
                        help="Path to client-guide-data.json (shape per references/schemas/client-guide.json)")
    parser.add_argument("--logo", required=True,
                        help="Path to CE logo PNG (download from Case Engine Branding folder; recommend 350x180 dark variant)")
    parser.add_argument("--firm", required=True, help="Law firm name (e.g., 'The May Firm')")
    parser.add_argument("--attorney", required=True, help="Attorney name being prepped (e.g., 'Robert May')")
    parser.add_argument("--practice-area", required=True,
                        help="Practice area name (Title Case, e.g., 'Car Accidents')")
    parser.add_argument("--episode-title", default=None,
                        help="Episode title - the MANDATORY prominent cover subtitle "
                             "(e.g., 'How to File a Car Accident Claim'). Falls back to "
                             "episode_title/episode_topic in the data JSON if omitted. "
                             "The cover renders this as the subtitle; the firm renders as a secondary line.")
    parser.add_argument("--scope", required=True, choices=["Topic Only", "Location", "Extension"],
                        help="Scope of the parent Client ROS this guide derives from")
    parser.add_argument("--location", default=None,
                        help="Jurisdictional location (e.g., 'California', 'CA - Long Beach'); required for Location/Extension")
    parser.add_argument("--output", required=True,
                        help='Output DOCX path - write it into the firm episode folder\'s "01 Strategy/" directory ALONGSIDE the ROS Template + Client ROS (both already there). Filename pattern "Client Guide - E{N} - {Episode Short Title} - {Location}.docx" (append " (Extension)" for an extension cell). Example: "<episode-folder>/01 Strategy/Client Guide - E2 - How to File a Car Accident Claim - CA.docx". The .md sibling is written alongside automatically. The DOCX is uploaded to Drive as a Google-Doc mimeType so Drive auto-converts to a clean branded Doc - never re-upload the .md with convert=true.')
    parser.add_argument("--run-date", default=None,
                        help="Run date in 'Month D, YYYY' format (defaults to today)")
    args = parser.parse_args()

    if args.scope in ("Location", "Extension") and not args.location:
        print("ERROR: --location is required when --scope is Location or Extension", file=sys.stderr)
        sys.exit(1)

    build(args)


if __name__ == "__main__":
    main()
