#!/usr/bin/env python3
"""Build the CE-branded ROS Template v2 DOCX and its paired markdown sibling.

The human-facing Google Doc MUST be this DOCX uploaded with
`mimeType: application/vnd.google-apps.document` so Drive auto-converts it.
NEVER upload the .md with convert=true - that leaks `[entity]{.underline}` as
visible text and has no cover page. The .md is uploaded as `text/markdown`,
unconverted, as the raw downstream-readable source.

Input:  ros-template-v2-data.json (validates against
        references/schema/ros-template-v2.json)
Output: a .docx and a .md, same basename.

Placeholder tokens ({{FIRM_NAME}}, {{YEARS_PRACTICING}}, ...) pass through
verbatim in both. Pandoc inline markers are translated to real Word runs in the
DOCX and stripped to plain text in the .md.

Dependencies: python-docx.

Usage:
  python3 build-ros-template-v2-docx.py \
      --data ros-template-v2-data.json \
      --logo /path/to/case-engine-logo.png \
      --output "E7: Truck Accidents // ROS Template v2 - GA - Savannah.docx"

Font is Roboto (see FONT). If the Case Engine Branding folder spec still says
Calibri, Roboto wins - flag the discrepancy when you spot it.
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# CE brand (per Case Engine Branding folder colors.md)
CE_BLUE = RGBColor(0x35, 0x73, 0xFF)
CE_DARK = RGBColor(0x0F, 0x17, 0x2A)
CE_GRAY = RGBColor(0x5B, 0x66, 0x76)
FONT = "Roboto"

# Everything this skill needs at runtime lives inside this skill. v2 is freestanding: it shares
# no module, helper or asset with pod-3A-ros-template or any other skill, so the directory can be
# zipped and handed over on its own. Nothing here may resolve to a path above SKILL_DIR.
SKILL_DIR = Path(__file__).resolve().parent.parent
DEFAULT_LOGO = SKILL_DIR / "assets" / "case-engine-logo.png"

# Schema order inside w:pPr - w:pBdr must precede all of these. From the ECMA-376 CT_PPr
# sequence, copied here rather than imported so the renderer does not reach into python-docx
# internals for it.
PPR_AFTER_PBDR = (
    "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap", "w:overflowPunct",
    "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
    "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
    "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
    "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)


# ---------- inline markup ----------

PANDOC_COLOR_RE = re.compile(r"\{\.color=[^}]+\}")
PANDOC_CLASS_RE = re.compile(r"\{\.[a-z][a-zA-Z0-9_-]*(?:=[^}]*)?\}")
PANDOC_INLINE_RE = re.compile(r"\[([^\]]+)\]\{\.(underline|smallcaps|mark)\}")

UTAG_RE = re.compile(r"<u>(.*?)</u>", re.S)


def _is_punct(ch):
    return not ch.isalnum() and not ch.isspace()


def _flanking(text, start, end):
    """Can the delimiter run text[start:end] open a span, close one, or both.

    CommonMark's left/right flanking test, cut down to what this renderer meets. It is what
    keeps a `**` that opens mid-sentence from being read as the close of the span it sits in.
    """
    prev = text[start - 1] if start > 0 else " "
    nxt = text[end] if end < len(text) else " "
    left = not nxt.isspace() and (not _is_punct(nxt) or prev.isspace() or _is_punct(prev))
    right = not prev.isspace() and (not _is_punct(prev) or nxt.isspace() or _is_punct(nxt))
    return left, right


def parse_inline(text, *, bold=False, italic=False, underline=False):
    """Return (chunk, bold, italic, underline) tuples for one paragraph of markup.

    A linear scan carrying a depth counter rather than a regex sweep, so a `**` opened inside
    an already open `**` span nests instead of closing it early. That case is real and it is
    the reason this was rewritten: a {{PLACEHOLDER}} wrapped in its own bold markers inside a
    fully bold Short-Form question used to close the outer run, and the token rendered unbold,
    which fails the placeholder gate. Depth is all that matters here - bold inside bold is
    still bold - so the parser never has to guess which marker pairs with which.

    Pandoc `[text]{.underline}` and `<u>` are consumed in place and carry whatever bold or
    italic is open around them, so a marker nested either way renders as one run holding both
    attributes and never leaks a literal marker.
    """
    runs = []
    if not text:
        return runs
    buf = []
    bd = 0
    itd = 0

    def flush():
        if not buf:
            return
        plain = PANDOC_COLOR_RE.sub("", PANDOC_CLASS_RE.sub("", "".join(buf)))
        del buf[:]
        if plain:
            runs.append((plain, bold or bd > 0, italic or itd > 0, underline))

    pos = 0
    n = len(text)
    while pos < n:
        ch = text[pos]
        if ch == "[":
            m = PANDOC_INLINE_RE.match(text, pos)
            if m:
                flush()
                runs.extend(parse_inline(m.group(1), bold=bold or bd > 0,
                                         italic=italic or itd > 0, underline=True))
                pos = m.end()
                continue
        elif ch == "<":
            m = UTAG_RE.match(text, pos)
            if m:
                flush()
                runs.extend(parse_inline(m.group(1), bold=bold or bd > 0,
                                         italic=italic or itd > 0, underline=True))
                pos = m.end()
                continue
        elif ch == "*":
            width = 2 if pos + 1 < n and text[pos + 1] == "*" else 1
            left, right = _flanking(text, pos, pos + width)
            if width == 2 and bd > 0 and right:
                flush()
                bd -= 1
                pos += 2
                continue
            if width == 2 and left:
                flush()
                bd += 1
                pos += 2
                continue
            if width == 1 and itd > 0 and right:
                flush()
                itd -= 1
                pos += 1
                continue
            if width == 1 and left:
                flush()
                itd += 1
                pos += 1
                continue
        buf.append(ch)
        pos += 1
    flush()
    return runs


def strip_pandoc(text):
    """Plain text for the .md sibling. Placeholder tokens pass through."""
    if text is None:
        return ""
    text = PANDOC_INLINE_RE.sub(r"\1", text)
    text = PANDOC_COLOR_RE.sub("", text)
    return PANDOC_CLASS_RE.sub("", text)


# ---------- docx primitives ----------

def set_run_style(run, *, size=11, bold=False, italic=False, underline=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def add_para(doc, text="", *, align=None, space_before=0, space_after=6,
             style=None, **run_kw):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_run_style(p.add_run(text), **run_kw)
    return p


def add_rich(doc, text, *, base_size=11, base_bold=False, base_italic=False,
             base_color=CE_DARK, space_before=0, space_after=6, style=None):
    """A paragraph whose inline **bold** / *italic* / [x]{.underline} become real runs."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for chunk, bo, it, un in parse_inline(text or ""):
        set_run_style(p.add_run(chunk), size=base_size, bold=bo or base_bold,
                      italic=it or base_italic, underline=un, color=base_color)
    return p


def add_h1(doc, text):
    return add_rich(doc, text, base_size=20, base_bold=True, base_color=CE_BLUE,
                    space_before=18, space_after=8)


def add_h2(doc, text):
    return add_rich(doc, text, base_size=16, base_bold=True, base_color=CE_DARK,
                    space_before=14, space_after=6)


def add_note(doc, text):
    """The italic gray direction lines. These SHIP - they are not internal notes."""
    return add_rich(doc, text, base_size=11, base_italic=True, base_color=CE_GRAY,
                    space_after=6)


def add_bullet(doc, text):
    return add_rich(doc, text, base_size=11, base_color=CE_DARK, space_after=3,
                    style="List Bullet")


# ---- S2 primitives ----
#
# S2 is scanned on the day, not read, so the block is deliberately tighter than the rest of the
# document: zero spacing on the questions and their bullets, small values on the location heading
# only. Compact is the format, not a rendering preference.

def add_sf_location(doc, text):
    return add_rich(doc, text, base_size=16, base_bold=True, base_color=CE_DARK,
                    space_before=8, space_after=2)


def add_sf_question(doc, text):
    """A Short-Form question. Bold in full, the Q label inside the bold, nothing else beside it."""
    return add_rich(doc, text, base_size=11, base_bold=True, base_color=CE_DARK,
                    space_before=0, space_after=0)


def add_sf_bullet(doc, text):
    """An attorney bullet. The label is underlined and NOT bold, so the question keeps the only
    bold weight in the block."""
    return add_rich(doc, text, base_size=11, base_color=CE_DARK, space_before=0, space_after=0,
                    style="List Bullet")


def add_rule(doc, *, space_before=12, space_after=12):
    """A real horizontal rule: an empty paragraph carrying a bottom border.

    Never a row of dash or underscore characters. Those are text - they wrap, they get read
    aloud, they leak into the .md as noise, and the QA gate greps for stray dash runs. The
    border element has to sit in schema order inside pPr (pBdr precedes spacing), hence the
    explicit successor list rather than an append.
    """
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = False
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "5B6676")
    pbdr.append(bottom)
    p._element.get_or_add_pPr().insert_element_before(pbdr, *PPR_AFTER_PBDR)
    return p


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _hdr_ftr_text(container, text, *, align=WD_ALIGN_PARAGRAPH.LEFT, size=9):
    p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.text = ""
    p.alignment = align
    set_run_style(p.add_run(text), size=size, color=CE_GRAY)


# ---------- document ----------

def scope_label(data):
    parts = [data.get("scope") or "Topic Only"]
    if data.get("location"):
        parts.append(data["location"])
    return " - ".join(parts)



# The frozen boilerplate strings (Editorial Guideline 8). The renderer asserts they are all
# present; it never invents one. Values live in ros-template-v2-data.json -> static and mirror
# references/statics.json -> strings. Eleven constants were retired 2026-08-17 and now live in
# references/statics.json -> retired; they are deliberately NOT required here.
# The outro's three spoken lines left the STATIC set 2026-08-18: Gabe required the outro to read
# unique per episode, so they are generated against the beats and banks in
# references/outro-banks.json. outro_note stays constant because it is host direction, not spoken.
STATIC_KEYS = (
    "welcome", "welcome_first",
)

# The three generated outro lines, in render order. Sign-off before reach-out is deliberate.
OUTRO_LINES = ("thanks", "signoff", "reach")

FOLLOWUP_NOTE = "Not read on air. For the interviewer and the attorney both."
FOLLOWUP_STATIC = ("Follow up when the opportunity presents itself, not on a schedule. "
                   "Let the answer finish first.")


def outro_lines(data):
    """The three generated outro lines, in render order.

    Not constants. Each is generated per episode against the required beats, invariants and banks
    in references/outro-banks.json, so the renderer reads them from the payload the same way it
    reads the prompt. It never carries a copy of any line; a missing one is a payload bug, not
    something to paper over with a hardcoded string. Any substitution on these strings is
    .replace() only - str.format would collapse the {{PLACEHOLDER}} braces they carry.
    """
    outro = data.get("outro") or {}
    missing = [k for k in OUTRO_LINES if not outro.get(k)]
    if missing:
        raise SystemExit(
            f"FAILED: outro block missing {len(missing)} generated line(s): {missing}. The outro's "
            "spoken lines stopped being constants 2026-08-18; generate them against "
            "references/outro-banks.json before render."
        )
    return [outro[k] for k in OUTRO_LINES]


# The line that sits under the S2 heading. Direction to the host, and it ships.
SF_NOTE = ("These question answers are designed to be clipped in short form and answered in 60 "
           "to 90 seconds.")


def sf_question(q, where):
    """The question text, with any bold markers of its own removed.

    The whole question renders bold, label included, so a token that arrives already wrapped in
    `**` is redundant at best. The parser nests safely now, but the markers are still stripped at
    the source so neither output carries a marker nobody asked for.
    """
    text = (q.get("q") or q.get("question") or "").strip()
    if not text:
        raise SystemExit(f"FAILED: {where} carries no question text.")
    return text.replace("**", "")


def sf_bullets(q):
    """The attorney bullets under a question, as `[Label]{.underline}: detail` strings.

    The label is underlined and never bold - the question keeps the only bold weight in the
    block. A bullet that already arrives as a formatted string passes through untouched.
    """
    out = []
    for b in q.get("bullets") or []:
        if isinstance(b, str):
            if b.strip():
                out.append(b.strip())
            continue
        label = (b.get("label") or "").strip()
        detail = (b.get("detail") or "").strip()
        if label:
            out.append(f"[{label}]{{.underline}}: {detail}".rstrip(": "))
        elif detail:
            out.append(detail)
    return out


def validate_attributes(data):
    """Gates AT-1 and AT-2 from references/attributes.md. Both are hard, not preferences.

    AT-1 is the gate that matters most: a question mark anywhere in the block means a catalog row
    reached the page uninverted, the attorney reads it aloud, and the block becomes the checklist
    interrogation v2 exists to replace. AT-2 caps the block at twelve, past which the attorney
    stops treating it as a shape and starts treating it as a list to get through.
    """
    attrs = data["segment_1"].get("attributes") or []
    if not 10 <= len(attrs) <= 12:
        raise SystemExit(
            f"FAILED: attribute block has {len(attrs)} bullet(s). Ten to twelve required (AT-2). "
            "See references/attributes.md."
        )
    for ai, a in enumerate(attrs, start=1):
        name = (a.get("name") or "").strip()
        detail = (a.get("detail") or "").strip()
        if "?" in name or "?" in detail:
            raise SystemExit(
                f"FAILED: attribute bullet {ai} ('{name}') carries a question mark. Zero question "
                "marks anywhere in the attribute block (AT-1). The block says what the attorney "
                "COVERS, never the question a client would ask. Rewrite it as response guidance; "
                "do not just delete the question mark."
            )


def validate_segment_2(data):
    """Ten questions per location, two to four bullets under each. Both are hard, not preferences."""
    for block in data["segment_2"]["locations"]:
        loc = block.get("location") or "(unnamed)"
        questions = block.get("questions") or []
        if not 8 <= len(questions) <= 10:
            raise SystemExit(
                f"FAILED: location '{loc}' has {len(questions)} questions. 8 to 10 required "
                "(Gabe 2026-08-24; was 15). Pick the best from the n-gram table's 20."
            )
        for qi, q in enumerate(questions, start=1):
            where = f"location '{loc}' Q{qi}"
            sf_question(q, where)
            n = len(sf_bullets(q))
            if not 2 <= n <= 4:
                raise SystemExit(
                    f"FAILED: {where} has {n} bullet(s). Two to four required, three the default. "
                    "Each is `[Label]{.underline}: detail`."
                )


def build_docx(data, logo, run_date):
    """Render the locked v2 shape. Mirrors scripts/reference-impl/push_v3.py exactly."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        _hdr_ftr_text(section.header,
                      f"Case Engine  |  Run of Show  |  {data['topic']}, {scope_label(data)}")
        _hdr_ftr_text(section.footer, "Case Engine  |  Confidential")

    S = data["static"]
    s1, s2 = data["segment_1"], data["segment_2"]
    cover = data.get("cover_page", {})

    # ---- cover page (spacer, logo at the 2nd paragraph, spacer) ----
    doc.add_paragraph()
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo:
        logo_p.add_run().add_picture(logo, width=Pt(180))
    doc.add_paragraph()

    add_para(doc, cover.get("title", "Run of Show"), align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=6, size=24, bold=True, color=CE_BLUE)
    add_para(doc, data["episode_title"], align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=18, bold=True, color=CE_DARK)
    doc.add_paragraph()
    add_para(doc, f"{data['topic']}  |  {scope_label(data)}", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=14, color=CE_DARK)
    doc.add_paragraph()
    add_para(doc, cover.get("prepared_by", "Prepared by Case Engine"),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, size=11, color=CE_DARK)
    add_page_break(doc)

    # ---- S1: Long-Form ----
    add_h1(doc, "S1: Long-Form (15-30m)")

    add_h2(doc, "Introduction")
    add_rich(doc, S["welcome"].replace("{topic_phrase}", s1["topic_phrase"]))
    add_rich(doc, s1.get("setup") or s1.get("cold_open") or "")
    # Line 3, the credential turn. Its own paragraph so the ask that follows lands as a turn
    # toward the attorney rather than as the tail of a long sentence. Not force-bold: only its
    # placeholders carry bold, so line 4 keeps the single bold prompt the gate asserts.
    if s1.get("credential"):
        add_rich(doc, s1["credential"])
    add_rich(doc, s1["prompt"], base_bold=True)
    add_h2(doc, "ATTORNEY RESPONSE")

    # The attribute bullets sit directly under the speaker tag. No heading above them and no
    # divider below. `Attributes to Hit` and the whole `Internal Notes (not read on air)` block
    # were retired 2026-08-17 per references/document-structure.md, and the QA removed-section
    # gate greps for both, so emitting either fails this skill's own QA step. Bullets, never a
    # numbered list (Gabe 2026-08-26).
    for a in s1["attributes"]:
        add_bullet(doc, f"**{a['name']}.** {a['detail']}")

    # ---- Outro: the end of S1 ----
    # S1 is a complete recording, intro through outro. The outro closes S1 rather than following
    # S2, which is a separate session in a different register.
    add_h2(doc, "Follow-ups")
    add_note(doc, FOLLOWUP_NOTE)
    add_bullet(doc, FOLLOWUP_STATIC)
    for f in (s1.get("follow_ups") or []):
        add_bullet(doc, f)

    add_h2(doc, "Outro")
    # Generated per episode. Thanks, sign-off, then the reach-out as a tag after the close - the
    # order is gated, and the sign-off is deliberately not last.
    for line in outro_lines(data):
        add_rich(doc, line)

    # ---- S2: Short-Form ----
    # No page break. S2 flows on from S1 behind a horizontal rule. Cover, S1 and the appendix are
    # the only sections that begin a page. pageBreakBefore is set explicitly rather than left
    # unset, because a paragraph otherwise inherits whatever sat at its index.
    add_rule(doc)
    add_h1(doc, "S2: Short-Form (60-90s)").paragraph_format.page_break_before = False
    add_note(doc, SF_NOTE)
    for block in s2["locations"]:
        loc = block["location"]
        add_sf_location(doc, f"Location: {loc}")
        for qi, q in enumerate(block["questions"], start=1):
            # The Q label sits INSIDE the bold so the whole line carries one weight. Nothing
            # else renders under a question: no time budget, no geo tag line, no source ref.
            add_sf_question(doc, f"**Q{qi}: {sf_question(q, loc)}**")
            for bullet in sf_bullets(q):
                add_sf_bullet(doc, bullet)

    # ---- appendix ----
    add_page_break(doc)
    add_h1(doc, "Appendix: Source Question Bank")
    add_note(doc, "The episode's N-Gram Table, verbatim. INTERNAL. In v2 this is reference rather "
                  "than script: Short-Form questions were rebuilt around search phrasing and "
                  "attributes, not lifted from here. Kept as the audit trail and the pull pool.")
    for row in data["appendix_question_bank"]:
        add_bullet(doc, f"**{row['n']}.** {row['question_text']}")

    return doc


# ---------- markdown sibling ----------

def build_markdown(data, run_date):
    """Same content shape as the DOCX, plain markdown. Pandoc markers stripped, tokens kept."""
    S = data["static"]
    s1, s2 = data["segment_1"], data["segment_2"]
    cover = data.get("cover_page", {})
    P = strip_pandoc
    out = []
    w = out.append

    w("# Run of Show")
    w("")
    w(f"**{data['episode_title']}**")
    w("")
    w(f"{data['topic']}  |  {scope_label(data)}")
    w("")
    w(cover.get("prepared_by", "Prepared by Case Engine"))
    w("")

    w("# S1: Long-Form (15-30m)")
    w("")
    w("## Introduction")
    w("")
    w(P(S["welcome"].replace("{topic_phrase}", s1["topic_phrase"])))
    w("")
    w(P(s1.get("setup") or s1.get("cold_open") or ""))
    w("")
    if s1.get("credential"):
        w(P(s1["credential"])); w("")
    w(f"**{P(s1['prompt']).replace('**', '')}**")
    w("")
    w("## ATTORNEY RESPONSE")
    w("")

    # No heading above the attribute bullets and no divider below them. Bullets, never a
    # numbered list (Gabe 2026-08-26).
    for a in s1["attributes"]:
        w(f"- **{a['name']}.** {P(a['detail'])}")
    w("")

    w("## Follow-ups")

    w("")

    w(f"*{FOLLOWUP_NOTE}*")

    w("")

    w(f"- {FOLLOWUP_STATIC}")

    for f in (s1.get("follow_ups") or []):

        w(f"- {f}")

    w("")

    w("## Outro")
    w("")
    w("")
    for line in outro_lines(data):
        w(P(line)); w("")

    # No page break between S1 and S2, a horizontal rule instead.
    w("---")
    w("")
    w("# S2: Short-Form (60-90s)")
    w("")
    w(f"*{SF_NOTE}*")
    w("")
    for block in s2["locations"]:
        loc = block["location"]
        w(f"## Location: {loc}")
        w("")
        for qi, q in enumerate(block["questions"], start=1):
            w(f"**Q{qi}: {P(sf_question(q, loc))}**")
            w("")
            for bullet in sf_bullets(q):
                w(f"- {P(bullet)}")
            w("")

    w("# Appendix: Source Question Bank")
    w("")
    w("*The episode's N-Gram Table, verbatim. INTERNAL. In v2 this is reference rather than script: "
      "Short-Form questions were rebuilt around search phrasing and attributes, not lifted from here. "
      "Kept as the audit trail and the pull pool.*")
    w("")
    for row in data["appendix_question_bank"]:
        w(f"- **{row['n']}.** {P(row['question_text'])}")
    w("")
    return "\n".join(out)


# ---------- entry point ----------

def build(args):
    data = json.loads(Path(args.data).read_text())
    if data.get("episode_format") != "v2-open-interview":
        raise SystemExit(
            "FAILED: episode_format is not 'v2-open-interview'. This renderer only builds v2 "
            "templates. A legacy template belongs to pod-3A-ros-template."
        )

    # STATIC verbatim gate - a regenerated constant must not reach the render.
    missing = [k for k in STATIC_KEYS if k not in data.get("static", {})]
    if missing:
        raise SystemExit(f"FAILED: static block missing {len(missing)} constant(s): {missing}")

    # The outro's spoken lines are generated rather than constant; fail before any render work.
    outro_lines(data)

    validate_attributes(data)
    validate_segment_2(data)

    logo = args.logo or (str(DEFAULT_LOGO) if DEFAULT_LOGO.is_file() else None)
    run_date = args.run_date or date.today().strftime("%B %d, %Y")
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    build_docx(data, logo, run_date).save(out)

    md_path = out.with_suffix(".md")
    md_path.write_text(build_markdown(data, run_date))

    locs = len(data["segment_2"]["locations"])
    print(f"DOCX: {out}")
    print(f"MD:   {md_path}")
    print(f"Data: {locs} location set(s) x 8-10 questions, "
          f"{len(data['appendix_question_bank'])} appendix bank rows, "
          f"{len(STATIC_KEYS)} static constants")
    if not logo:
        print(f"WARN: no --logo given and no bundled logo at {DEFAULT_LOGO}, cover page "
              "rendered without the CE logo")


def main():
    p = argparse.ArgumentParser(description="Build the CE-branded ROS Template v2 DOCX + .md")
    p.add_argument("--data", required=True, help="Path to ros-template-v2-data.json")
    p.add_argument("--logo", default=None,
                   help="Path to the CE logo image for the cover page. Defaults to the bundled "
                        "assets/case-engine-logo.png when that file is present.")
    p.add_argument("--output", required=True, help="Output .docx path (the .md is written alongside)")
    p.add_argument("--run-date", default=None, help="Cover-page date; defaults to today")
    build(p.parse_args())


if __name__ == "__main__":
    sys.exit(main())
