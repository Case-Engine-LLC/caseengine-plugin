#!/usr/bin/env python3
"""
What: Builds a CE-branded Entity Map DOCX from an entity-map.json + vector-space PNG.
Input: entity-map.json (filename per the canonical Podcast Drive doc:
       https://docs.google.com/document/d/1YhybGpp9DIqmV56P6OOHIQe7A6RxvjQGHLHKcUM0JmU)
       + Entity Vector Space.png (produced by the sibling entity-vector-space.py script)
       + a CE logo image (download from the canonical Case Engine Branding folder:
       https://drive.google.com/drive/folders/1OulNcg6hZ6caFD0Xc-W50i7ougC_y0qo)
Output: Entity Map.docx with cover page + headers/footers + tier tables + cluster
        architecture + bridge entities + embedded vector-space chart.
Re-run: Safe - overwrites the output DOCX. Drive's auto-conversion preserves the same fileId
        when the calling skill uses files.update for re-uploads (per the Push to Drive doc).

Convention sync: section order, branding spec, and the canonical filename live in the
Podcast Drive doc + the Case Engine Branding folder. If those change, update this script
to match. The canonical sources are the single source of truth.

Dependencies: python-docx (cowork runtime has it). No optional deps.

Usage:
  python3 build-entity-map-docx.py \\
      --json /path/to/entity-map.json \\
      --chart /path/to/Entity-Vector-Space.png \\
      --logo /path/to/ce-logo.png \\
      --practice-area "Car Accidents" \\
      --scope "Topic Only" \\
      --output /path/to/Entity-Map.docx \\
      [--location "CA - Long Beach"]   (omit at Topic Only scope)
      [--run-date "April 27, 2026"]    (defaults to today)
"""

import argparse
import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# CE brand (per Case Engine Branding folder colors.md)
CE_BLUE = RGBColor(0x35, 0x73, 0xFF)
CE_DARK = RGBColor(0x0F, 0x17, 0x2A)
CE_GRAY = RGBColor(0x5B, 0x66, 0x76)
FONT = "Calibri"


def set_run_style(run, *, font=FONT, size=11, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)


def add_para(doc, text="", *, align=None, space_before=0, space_after=6, **run_kw):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_style(run, **run_kw)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_header(section, text):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_style(run, size=9, italic=True, color=CE_GRAY)


def set_footer(section, text_left):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text_left + "\t\t")
    set_run_style(run, size=9, color=CE_GRAY)
    run2 = p.add_run("Page ")
    set_run_style(run2, size=9, color=CE_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_style(run, size=16, bold=True, color=CE_BLUE)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_style(run, size=13, bold=True, color=CE_DARK)
    return p


def add_table(doc, headers, rows, *, bold_rows=None):
    bold_rows = bold_rows or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_style(run, size=10, bold=True, color=CE_BLUE)
    for r_idx, row_vals in enumerate(rows):
        row_cells = table.rows[1 + r_idx].cells
        is_bold = r_idx in bold_rows
        for i, v in enumerate(row_vals):
            row_cells[i].text = ""
            run = row_cells[i].paragraphs[0].add_run(str(v))
            set_run_style(run, size=10, bold=is_bold, color=CE_DARK)
    return table


def titlecase_type(t):
    return " ".join(w.capitalize() for w in t.replace("_", " ").split())


def generate_insights(entities, clusters, bridges, tier_1, tier_2, tier_3):
    """Auto-generate landscape observations from the entity map data.

    Returns a list of insight strings. Heuristics:
    - Top entity types (where the domain leans)
    - Largest cluster (what the domain centers on)
    - Highest-connectivity bridge (the central node)
    - Tier-balance shape (top-heavy / long-tail / balanced)
    - Mean vector strength per tier (consensus signal)
    """
    insights = []
    total = len(entities)

    # Top entity types
    type_counts = {}
    for e in entities:
        type_counts[e["type"]] = type_counts.get(e["type"], 0) + 1
    top_types = sorted(type_counts.items(), key=lambda x: -x[1])[:3]
    if top_types and total > 0:
        top_pct = sum(c for _, c in top_types) / total * 100
        type_labels = ", ".join(f"{titlecase_type(t)} ({c})" for t, c in top_types)
        insights.append(
            f"**Domain leaning:** top 3 entity types are {type_labels} - {top_pct:.0f}% of all entities. "
            f"The domain is most heavily {titlecase_type(top_types[0][0])}."
        )

    # Largest cluster
    if clusters:
        largest = max(clusters.items(), key=lambda kv: len(kv[1].get("entities", [])))
        l_name = largest[1].get("name", largest[0])
        l_count = len(largest[1].get("entities", []))
        insights.append(
            f"**Largest cluster:** {l_name} ({l_count} entities) - the gravitational center of this practice area."
        )

    # Highest-connectivity bridge
    if bridges:
        top_bridge = max(bridges, key=lambda b: len(b.get("clusters_connected", [])))
        b_name = top_bridge["name"]
        b_clusters = top_bridge.get("clusters_connected", [])
        b_reason = top_bridge.get("reason", "")
        cls_names = " + ".join(clusters.get(c, {}).get("name", c) for c in b_clusters)
        line = f"**Central bridge:** {b_name} spans {len(b_clusters)} clusters ({cls_names})."
        if b_reason:
            line += f" {b_reason}"
        insights.append(line)

    # Tier-balance shape
    if total > 0:
        t1_pct = len(tier_1) / total * 100
        t3_pct = len(tier_3) / total * 100
        if t1_pct > 30:
            shape = f"**Top-heavy distribution** ({t1_pct:.0f}% Tier 1) - strong consensus on a small set of core entities; downstream content can lean hard on Tier 1."
        elif t3_pct > 40:
            shape = f"**Long-tail distribution** ({t3_pct:.0f}% Tier 3) - broad practice with many specialized terms; expect niche episodes drawing from supporting entities."
        else:
            shape = f"**Balanced distribution** (T1 {t1_pct:.0f}% / T2 {len(tier_2)/total*100:.0f}% / T3 {t3_pct:.0f}%) - healthy spread across core, major, and supporting tiers."
        insights.append(shape)

    # Mean vector strength per tier (consensus signal)
    def mean_vs(tier):
        vs = [e["vector_strength"] for e in tier]
        return sum(vs) / len(vs) if vs else 0
    if tier_1 or tier_2:
        line = (
            f"**Score consensus:** "
            f"Tier 1 mean vector strength {mean_vs(tier_1):.2f}, "
            f"Tier 2 {mean_vs(tier_2):.2f}, "
            f"Tier 3 {mean_vs(tier_3):.2f}."
        )
        if tier_1 and mean_vs(tier_1) > 0.88:
            line += " High Tier 1 consensus - the core is sharp and well-defined."
        elif tier_1 and mean_vs(tier_1) < 0.83:
            line += " Tier 1 consensus is soft - the core entities are close to the threshold; may benefit from re-scoring."
        insights.append(line)

    # Bridge tier composition
    if bridges and entities:
        bridge_tiers = []
        for b in bridges:
            ent = next((e for e in entities if e["id"] == b["id"]), None)
            if ent:
                bridge_tiers.append(ent["tier"])
        t1_bridges = sum(1 for t in bridge_tiers if t == 1)
        if t1_bridges == len(bridges):
            insights.append(
                f"**Bridge composition:** all {len(bridges)} bridges are Tier 1 - the topical graph reconverges on the most prominent entities."
            )
        elif t1_bridges >= len(bridges) // 2:
            insights.append(
                f"**Bridge composition:** majority Tier 1 ({t1_bridges}/{len(bridges)}) - core entities carry the cross-cluster connections."
            )
        else:
            insights.append(
                f"**Bridge composition:** mixed tiers ({t1_bridges} T1 of {len(bridges)} total) - cross-cluster connections come from across the tier hierarchy."
            )

    return insights


def generate_plain_translation(entities, clusters, bridges, tier_1, tier_2, tier_3, practice_area):
    """Translate the technical insights into client-facing plain language.

    Returns a list of plain-language strings with markdown ** for emphasis.
    Uses the same data signals as generate_insights but writes for a non-technical reader.
    """
    lines = []
    total = len(entities)

    # Domain leaning - translate to "this practice area is fundamentally about X"
    type_counts = {}
    for e in entities:
        type_counts[e["type"]] = type_counts.get(e["type"], 0) + 1
    top_types = sorted(type_counts.items(), key=lambda x: -x[1])[:3]
    if top_types and total > 0:
        top_type, top_count = top_types[0]
        top_type_pct = top_count / total * 100
        top_label = titlecase_type(top_type).lower()
        if top_type_pct > 25:
            lines.append(
                f"This practice area is fundamentally about **{top_label}**. "
                f"Roughly {top_type_pct:.0f}% of all entities fall in that bucket - "
                f"unusually concentrated. Content has to engage that subject directly; you can't sidestep it."
            )
        else:
            lines.append(
                f"The practice area is balanced across multiple subject types. "
                f"No single category dominates - top three are {', '.join(titlecase_type(t).lower() for t,_ in top_types)}. "
                f"Content can range broadly without feeling off-topic."
            )

    # Largest cluster - what the conversation centers on
    if clusters:
        largest = max(clusters.items(), key=lambda kv: len(kv[1].get("entities", [])))
        l_name = largest[1].get("name", largest[0])
        l_count = len(largest[1].get("entities", []))
        lines.append(
            f"The biggest theme in this domain is **{l_name}** - "
            f"more entities than any other cluster ({l_count}). "
            f"Audiences expect a real conversation about this; if your content skips it, you're missing the gravitational center."
        )

    # Central bridge
    if bridges:
        top_bridge = max(bridges, key=lambda b: len(b.get("clusters_connected", [])))
        b_name = top_bridge["name"]
        b_clusters_count = len(top_bridge.get("clusters_connected", []))
        cls_names = [clusters.get(c, {}).get("name", c) for c in top_bridge.get("clusters_connected", [])]
        if cls_names:
            connected = ", ".join(cls_names[:-1]) + f" and {cls_names[-1]}" if len(cls_names) > 1 else cls_names[0]
            lines.append(
                f"**{b_name}** is the connective tissue across the whole practice area. "
                f"It ties together {connected}. "
                f"If your content explains it well, you're educating AND linking related topics in one move."
            )

    # Tier-balance shape - what content strategy that implies
    if total > 0:
        t1_pct = len(tier_1) / total * 100
        t3_pct = len(tier_3) / total * 100
        if t1_pct > 30:
            lines.append(
                "**The core ideas are tightly defined.** There's strong consensus on what matters most. "
                "That means content can lean confidently on the same handful of central concepts - "
                "you don't need to invent novel angles, you need to explain the well-known ones better than competitors."
            )
        elif t3_pct > 40:
            lines.append(
                "**The practice area is broad** - lots of specialized ideas without a tight core. "
                "Expect niche episodes pulling from the supporting tier. "
                "Strategy: cover the core well, then differentiate with depth on a couple of niches."
            )
        else:
            lines.append(
                "**The landscape is balanced.** Core, major, and supporting ideas all carry weight. "
                "Content can range across the full hierarchy without feeling top-heavy or scattered."
            )

    # Bridge composition (whether the connectors ARE the core)
    if bridges and entities:
        bridge_tiers = []
        for b in bridges:
            ent = next((e for e in entities if e["id"] == b["id"]), None)
            if ent:
                bridge_tiers.append(ent["tier"])
        t1_bridges = sum(1 for t in bridge_tiers if t == 1)
        if t1_bridges == len(bridges):
            lines.append(
                "**The connectors are the most prominent ideas.** Coverage of the bridge entities does double duty - "
                "explains the central concepts AND ties related topics together. Highest-leverage content lives here."
            )

    return lines


def build(args):
    json_path = Path(args.json)
    if not json_path.exists():
        print(f"ERROR: {json_path} not found", file=sys.stderr)
        sys.exit(1)
    data = json.loads(json_path.read_text())

    practice_area = args.practice_area
    scope = args.scope
    location = args.location
    run_date = args.run_date or date.today().strftime("%B %-d, %Y")

    entities = data.get("entities", [])
    clusters = data.get("connection_graph", {}).get("clusters", {})
    bridges = data.get("connection_graph", {}).get("bridge_entities", [])
    bridge_names = {b["name"] for b in bridges}

    tier_1 = sorted([e for e in entities if e["tier"] == 1], key=lambda e: -e["vector_strength"])
    tier_2 = sorted([e for e in entities if e["tier"] == 2], key=lambda e: -e["vector_strength"])
    tier_3 = sorted([e for e in entities if e["tier"] == 3], key=lambda e: -e["vector_strength"])

    scope_label = f"{scope}" + (f" - {location}" if location else "")
    header_text = f"Case Engine  |  Entity Research  |  {practice_area}, {scope_label}"

    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        set_header(section, header_text)
        set_footer(section, "Case Engine  |  Confidential")

    # ---- Cover page ----
    for _ in range(3):
        doc.add_paragraph()

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(args.logo, width=Inches(2.4))

    add_para(doc, "Entity Research", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=6, size=36, bold=True, color=CE_BLUE)
    add_para(doc, practice_area, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=22, bold=True, color=CE_DARK)
    add_para(doc, scope_label + (" - Foundation Map" if scope == "Topic Only" else ""),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24, size=14, color=CE_DARK)
    add_para(doc, "Prepared by Case Engine", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=2, size=11, italic=True, color=CE_GRAY)
    add_para(doc, run_date, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=0, size=11, color=CE_GRAY)

    add_page_break(doc)

    # ---- Executive Summary ----
    add_h2(doc, "Executive Summary")
    intro = (
        f"This map captures the entity universe for {practice_area} at {scope_label} scope. "
        + ("Foundation map - jurisdiction-agnostic; every Location and Extension build inherits from this."
           if scope == "Topic Only"
           else "Localized cascade - inherits jurisdiction-neutral entities from Topic Only, then adds locally-strong entities that score on the vector formula.")
    )
    add_para(doc, intro, size=11, color=CE_DARK, space_after=10)

    add_h3(doc, "Counts")
    for line in [
        f"{len(entities)} entities (target: 40-50)",
        f"{len(clusters)} clusters (target: 8-15)",
        f"{len(bridges)} bridge entities (target: 4-6)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(line)
        set_run_style(run, size=11, color=CE_DARK)

    add_h3(doc, "Tier distribution")
    for line in [
        f"Tier 1 (core, vector strength >= 0.80): {len(tier_1)} entities",
        f"Tier 2 (major, 0.60-0.79): {len(tier_2)} entities",
        f"Tier 3 (supporting, 0.40-0.59): {len(tier_3)} entities",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(line)
        set_run_style(run, size=11, color=CE_DARK)

    add_h3(doc, "Localization")
    if scope == "Topic Only":
        loc_lines = [
            "Coverage: 0% jurisdictional-named (expected at Topic Only)",
            "Supplement: not triggered, not applicable at this scope",
        ]
    else:
        loc = data.get("localization", {})
        cov = loc.get("coverage_pct", "n/a")
        sup_status = loc.get("supplement_status", "not-triggered")
        loc_lines = [
            f"Coverage: {cov}% jurisdictional-named",
            f"Supplement: {sup_status}",
        ]
    for line in loc_lines:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(line)
        set_run_style(run, size=11, color=CE_DARK)

    # ---- Learnings & Insights (auto-generated from data) ----
    add_h3(doc, "Learnings & Insights")
    insights = generate_insights(entities, clusters, bridges, tier_1, tier_2, tier_3)
    for insight in insights:
        p = doc.add_paragraph(style="List Bullet")
        # Render bold + plain text inline by splitting on **
        parts = insight.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            set_run_style(run, size=11, bold=(i % 2 == 1), color=CE_DARK)

    # ---- What does this mean? (client-facing translation of the insights) ----
    add_h3(doc, "What does this mean?")
    plain_lines = generate_plain_translation(entities, clusters, bridges, tier_1, tier_2, tier_3, practice_area)
    for line in plain_lines:
        p = doc.add_paragraph(style="List Bullet")
        parts = line.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            set_run_style(run, size=11, bold=(i % 2 == 1), color=CE_DARK)

    # Page break - Vector Space Visualization starts fresh on page 3
    add_page_break(doc)

    # ---- Vector Space Visualization (page 3, after exec summary + insights) ----
    add_h2(doc, "Vector Space Visualization")
    add_para(doc,
             "Entities plotted radially - cluster determines angle, vector strength determines distance from center. Bridges highlighted with gold border.",
             size=11, italic=True, color=CE_GRAY, space_after=8)
    chart_path = Path(args.chart) if args.chart else None
    if chart_path and chart_path.exists():
        chart_p = doc.add_paragraph()
        chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_p.add_run().add_picture(str(chart_path), width=Inches(6.0))
    else:
        add_para(doc, "[Vector space chart not available - run entity-vector-space.py first]",
                 size=10, italic=True, color=CE_GRAY)

    # ---- Tier tables ----
    headers = ["Entity", "Type", "Vec Str", "Prom", "Rel", "Pop", "Cluster", "Bridge"]

    def build_tier_rows(tier_entities):
        rows = []
        bold_idx = set()
        for i, e in enumerate(tier_entities):
            cluster_display = clusters.get(e["cluster"], {}).get("name", e["cluster"])
            is_bridge = e["name"] in bridge_names
            if is_bridge:
                bold_idx.add(i)
            rows.append([
                e["name"],
                titlecase_type(e["type"]),
                f"{e['vector_strength']:.3f}",
                f"{e['prominence']:.2f}",
                f"{e['relatedness']:.2f}",
                f"{e['popularity']:.2f}",
                cluster_display,
                "*" if is_bridge else "-",
            ])
        return rows, bold_idx

    add_h2(doc, "Tier 1 - Core Entities")
    add_para(doc, "Central to the practice area; appears in nearly every episode at every scope.",
             size=11, color=CE_DARK, space_after=8)
    rows, bold = build_tier_rows(tier_1)
    bold = bold | set(range(len(tier_1)))
    add_table(doc, headers, rows, bold_rows=bold)

    add_h2(doc, "Tier 2 - Major Entities")
    add_para(doc, "Important; appears in most episodes.",
             size=11, color=CE_DARK, space_after=8)
    rows, bold = build_tier_rows(tier_2)
    add_table(doc, headers, rows, bold_rows=bold)

    add_h2(doc, "Tier 3 - Supporting Entities")
    add_para(doc, "Niche / specialized; appears where relevant.",
             size=11, color=CE_DARK, space_after=8)
    rows, bold = build_tier_rows(tier_3)
    add_table(doc, headers, rows, bold_rows=bold)

    # ---- Cluster Architecture ----
    add_h2(doc, "Cluster Architecture")
    add_para(doc,
             f"{len(clusters)} contextual layers slice this practice area. Each cluster represents one way the domain is sliced; together they form the topical map.",
             size=11, color=CE_DARK, space_after=8)

    for cluster_key, cluster_data in clusters.items():
        cluster_name = cluster_data.get("name", cluster_key)
        cluster_desc = cluster_data.get("contextual_layer", "")
        cluster_entities_ids = cluster_data.get("entities", [])

        add_h3(doc, cluster_name)
        if cluster_desc:
            add_para(doc, cluster_desc, size=11, italic=True, color=CE_GRAY, space_after=4)

        for eid in cluster_entities_ids:
            ent = next((e for e in entities if e["id"] == eid), None)
            if not ent:
                continue
            tier_label = f"Tier {ent['tier']}"
            is_bridge = ent["name"] in bridge_names
            tag = f"{tier_label}, bridge" if is_bridge else tier_label
            p = doc.add_paragraph(style="List Bullet")
            r1 = p.add_run(ent["name"])
            set_run_style(r1, size=11, bold=True, color=CE_DARK)
            r2 = p.add_run(f"  ({tag})")
            set_run_style(r2, size=11, color=CE_DARK)

    # ---- Bridge Entities ----
    add_h2(doc, "Bridge Entities")
    add_para(doc,
             "Bridges connect multiple clusters and carry the highest authority value because they are where the topical graph reconverges.",
             size=11, color=CE_DARK, space_after=8)
    bridge_rows = []
    for b in bridges:
        ent = next((e for e in entities if e["id"] == b["id"]), None)
        tier_label = f"T{ent['tier']}" if ent else "?"
        cls_names = " + ".join(
            clusters.get(c, {}).get("name", c) for c in b.get("clusters_connected", [])
        )
        bridge_rows.append([b["name"], tier_label, cls_names, len(b.get("clusters_connected", []))])
    add_table(doc, ["Bridge", "Tier", "Connects", "Connections"], bridge_rows,
              bold_rows=set(range(len(bridge_rows))))

    # ---- Localization Summary ----
    add_h2(doc, "Localization Summary")
    if scope == "Topic Only":
        add_para(doc,
                 "This map is at Topic Only scope - jurisdiction-agnostic by design. No location-specific entities are forced. When this map is inherited by a Location or Extension build, locally-strong entities (named police departments, local hospitals, state-specific statutes/forms) enter via the localization-coverage evaluation step.",
                 size=11, color=CE_DARK, space_after=4)
        add_para(doc, "Coverage at this scope: 0% jurisdictional-named (expected). Supplement: not triggered.",
                 size=11, italic=True, color=CE_GRAY, space_after=12)
    else:
        loc = data.get("localization", {})
        cov = loc.get("coverage_pct", "n/a")
        sup_status = loc.get("supplement_status", "not-triggered")
        add_para(doc,
                 f"This map is at {scope_label}. Coverage: {cov}% jurisdictional-named. Supplement status: {sup_status}.",
                 size=11, color=CE_DARK, space_after=12)

    # ---- Inheritance Notes ----
    add_h2(doc, "Inheritance Notes")
    if scope == "Topic Only":
        add_para(doc,
                 "This Topic Only map is the foundation that every Location and Extension cascade reads first. When building below this scope:",
                 size=11, color=CE_DARK, space_after=8)
        for line in [
            "Carry forward Tier 1 + Tier 2 entities that are jurisdiction-neutral",
            "Add entities unique to the jurisdiction (named police departments, local hospitals, state-specific statutes/forms)",
            "Don't re-score universal entities",
        ]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line)
            set_run_style(run, size=11, color=CE_DARK)
    else:
        add_para(doc,
                 f"This {scope} map inherits from the Topic Only foundation for {practice_area}. Universal entities are carried over without re-scoring; locally-strong entities are added.",
                 size=11, color=CE_DARK, space_after=8)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved DOCX: {output_path}")

    # Also emit a sibling .md alongside the .docx with the same content shape.
    # The .md is the raw markdown deliverable Drive will upload as `text/markdown` (no auto-conversion).
    # The .docx is the Drive-rendered Google Doc sibling. Both live side-by-side in the scope folder.
    md_path = output_path.with_suffix(".md")
    md_text = build_markdown(
        practice_area=practice_area, scope=scope, location=location, run_date=run_date,
        entities=entities, clusters=clusters, bridges=bridges, bridge_names=bridge_names,
        tier_1=tier_1, tier_2=tier_2, tier_3=tier_3,
        chart_relpath=("visuals/" + Path(args.chart).name) if args.chart else None,
    )
    md_path.write_text(md_text)
    print(f"Saved MD:   {md_path}")
    print(f"Data: {len(entities)} entities, {len(clusters)} clusters, {len(bridges)} bridges")


def build_markdown(*, practice_area, scope, location, run_date, entities, clusters, bridges,
                   bridge_names, tier_1, tier_2, tier_3, chart_relpath):
    """Build the raw .md deliverable - same content shape as the DOCX, plain markdown.

    The .md ships alongside the .docx in Drive (uploaded as text/markdown without conversion).
    Downstream skills that prefer raw markdown can read the .md directly; humans can open
    the Google Doc sibling.
    """
    scope_label = f"{scope}" + (f" - {location}" if location else "")
    lines = []

    lines.append(f"# Entity Research - {practice_area} ({scope_label})")
    lines.append("")
    lines.append(f"_Prepared by Case Engine - {run_date}_")
    lines.append("")
    lines.append("---")
    lines.append("")

    lines.append("## Executive Summary")
    lines.append("")
    if scope == "Topic Only":
        intro = (f"This map captures the entity universe for {practice_area} at {scope_label} scope. "
                 "Foundation map - jurisdiction-agnostic; every Location and Extension build inherits from this.")
    else:
        intro = (f"This map captures the entity universe for {practice_area} at {scope_label} scope. "
                 "Localized cascade - inherits jurisdiction-neutral entities from Topic Only, then adds locally-strong entities that score on the vector formula.")
    lines.append(intro)
    lines.append("")

    lines.append("### Counts")
    lines.append("")
    lines.append(f"- {len(entities)} entities (target: 40-50)")
    lines.append(f"- {len(clusters)} clusters (target: 8-15)")
    lines.append(f"- {len(bridges)} bridge entities (target: 4-6)")
    lines.append("")

    lines.append("### Tier distribution")
    lines.append("")
    lines.append(f"- Tier 1 (core, vector strength >= 0.80): {len(tier_1)} entities")
    lines.append(f"- Tier 2 (major, 0.60-0.79): {len(tier_2)} entities")
    lines.append(f"- Tier 3 (supporting, 0.40-0.59): {len(tier_3)} entities")
    lines.append("")

    lines.append("### Localization")
    lines.append("")
    if scope == "Topic Only":
        lines.append("- Coverage: 0% jurisdictional-named (expected at Topic Only)")
        lines.append("- Supplement: not triggered, not applicable at this scope")
    else:
        lines.append("- Coverage: see metadata for jurisdictional-coverage % at this scope")
        lines.append("- Supplement: see `localization_supplement` block in entity-map.json (null unless triggered)")
    lines.append("")

    insights = generate_insights(entities, clusters, bridges, tier_1, tier_2, tier_3)
    if insights:
        lines.append("### Learnings & Insights")
        lines.append("")
        for insight in insights:
            lines.append(f"- {insight}")
        lines.append("")

    plain_lines = generate_plain_translation(entities, clusters, bridges, tier_1, tier_2, tier_3, practice_area)
    if plain_lines:
        lines.append("### What does this mean?")
        lines.append("")
        for line in plain_lines:
            lines.append(f"- {line}")
        lines.append("")

    lines.append("## Vector Space Visualization")
    lines.append("")
    lines.append("Entities plotted radially - cluster determines angle, vector strength determines distance from center. Bridges highlighted with gold border.")
    lines.append("")
    if chart_relpath:
        lines.append(f"![Vector space chart]({chart_relpath})")
        lines.append("")

    def render_tier_table(tier_entities, name, intro):
        if not tier_entities:
            return
        lines.append(f"## {name}")
        lines.append("")
        lines.append(intro)
        lines.append("")
        lines.append("| Entity | Type | Vec Str | Prom | Rel | Pop | Cluster | Bridge |")
        lines.append("|---|---|---|---|---|---|---|---|")
        for e in tier_entities:
            cluster_display = clusters.get(e["cluster"], {}).get("name", e["cluster"])
            is_bridge = e["name"] in bridge_names
            name_cell = f"**{e['name']}**" if is_bridge else e["name"]
            bridge_cell = "*" if is_bridge else "-"
            lines.append(
                f"| {name_cell} | {titlecase_type(e['type'])} | "
                f"{e['vector_strength']:.3f} | {e['prominence']:.2f} | "
                f"{e['relatedness']:.2f} | {e['popularity']:.2f} | "
                f"{cluster_display} | {bridge_cell} |"
            )
        lines.append("")

    render_tier_table(tier_1, "Tier 1 - Core Entities",
                      "Central to the practice area; appears in nearly every episode at every scope.")
    render_tier_table(tier_2, "Tier 2 - Major Entities",
                      "Important; appears in most episodes.")
    render_tier_table(tier_3, "Tier 3 - Supporting Entities",
                      "Niche / specialized; appears where relevant.")

    lines.append("## Cluster Architecture")
    lines.append("")
    lines.append(f"{len(clusters)} contextual layers slice this practice area. Each cluster represents one way the domain is sliced; together they form the topical map.")
    lines.append("")
    for cluster_key, cluster_data in clusters.items():
        cluster_name = cluster_data.get("name", cluster_key)
        cluster_desc = cluster_data.get("contextual_layer", "")
        lines.append(f"### {cluster_name}")
        lines.append("")
        if cluster_desc:
            lines.append(f"_{cluster_desc}_")
            lines.append("")
        for eid in cluster_data.get("entities", []):
            ent = next((e for e in entities if e["id"] == eid), None)
            if not ent:
                continue
            tier_label = f"Tier {ent['tier']}"
            is_bridge = ent["name"] in bridge_names
            tag = f"{tier_label}, bridge" if is_bridge else tier_label
            lines.append(f"- **{ent['name']}** ({tag})")
        lines.append("")

    lines.append("## Bridge Entities")
    lines.append("")
    lines.append("Bridges connect multiple clusters and carry the highest authority value because they are where the topical graph reconverges.")
    lines.append("")
    lines.append("| Bridge | Tier | Connects | Connections |")
    lines.append("|---|---|---|---|")
    for b in bridges:
        ent = next((e for e in entities if e["id"] == b["id"]), None)
        tier_label = f"T{ent['tier']}" if ent else "?"
        cls_names = " + ".join(
            clusters.get(c, {}).get("name", c) for c in b.get("clusters_connected", [])
        )
        lines.append(f"| **{b['name']}** | {tier_label} | {cls_names} | {len(b.get('clusters_connected', []))} |")
    lines.append("")

    lines.append("## Localization Summary")
    lines.append("")
    if scope == "Topic Only":
        lines.append("This map is at Topic Only scope - jurisdiction-agnostic by design. No location-specific entities are forced. When this map is inherited by a Location or Extension build, locally-strong entities (named police departments, local hospitals, state-specific statutes/forms) enter via the localization-coverage evaluation step.")
        lines.append("")
        lines.append("_Coverage at this scope: 0% jurisdictional-named (expected). Supplement: not triggered._")
    else:
        lines.append(f"This map is at {scope_label}. Coverage and supplement status logged in `entity-map.json` metadata.")
    lines.append("")

    lines.append("## Inheritance Notes")
    lines.append("")
    if scope == "Topic Only":
        lines.append("This Topic Only map is the foundation that every Location and Extension cascade reads first. When building below this scope:")
        lines.append("")
        lines.append("- Carry forward Tier 1 + Tier 2 entities that are jurisdiction-neutral")
        lines.append("- Add entities unique to the jurisdiction (named police departments, local hospitals, state-specific statutes/forms)")
        lines.append("- Don't re-score universal entities")
    else:
        lines.append(f"This {scope} map inherits from the Topic Only foundation for {practice_area}. Universal entities are carried over without re-scoring; locally-strong entities are added.")
    lines.append("")

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="Build CE-branded Entity Map DOCX from entity-map.json")
    parser.add_argument("--json", required=True, help="Path to entity-map.json")
    parser.add_argument("--chart", default=None, help="Path to Entity Vector Space.png (optional - degrades gracefully if absent)")
    parser.add_argument("--logo", required=True, help="Path to CE logo PNG (download from Case Engine Branding folder; recommend 350x180 dark variant)")
    parser.add_argument("--practice-area", required=True, help="Practice area name (Title Case, e.g. 'Car Accidents')")
    parser.add_argument("--scope", required=True, choices=["Topic Only", "Location", "Extension"],
                        help="Scope of this map")
    parser.add_argument("--location", default=None,
                        help="State-prefixed jurisdictional folder name (e.g. 'CA - Long Beach'); required for Location/Extension")
    parser.add_argument("--output", required=True, help="Output DOCX path (e.g. <scope-folder>/Entity Map.docx)")
    parser.add_argument("--run-date", default=None, help="Run date in 'Month D, YYYY' format (defaults to today)")
    args = parser.parse_args()

    if args.scope in ("Location", "Extension") and not args.location:
        print("ERROR: --location is required when --scope is Location or Extension", file=sys.stderr)
        sys.exit(1)

    build(args)


if __name__ == "__main__":
    main()
