#!/usr/bin/env python3
"""Build the CE-branded Client ROS v2 DOCX and its paired markdown sibling.

The Client ROS v2 is the POPULATED copy of a v2 ROS Template - every one of the
twelve {{PLACEHOLDER}} tokens resolved to a firm's real values, the internal
`# Appendix: Source Question Bank` STRIPPED, and everything else byte-identical
to the template shape. This renderer is the populate-side mirror of
pod-3A-ros-template-v2/scripts/build-ros-template-v2-docx.py: same primitives,
same locked shape, same gates - plus the populate gates (zero leftover tokens,
statics match the template constants after token resolution).

The human-facing Google Doc MUST be this DOCX uploaded with
`mimeType: application/vnd.google-apps.document` so Drive auto-converts it.
NEVER upload the .md with convert=true - that leaks `[entity]{.underline}` as
visible text and has no cover page. The .md is uploaded as `text/markdown`,
unconverted, as the raw machine-readable source of record for the episode.

Input:  client-ros-v2-data.json (validates against
        references/schema/client-ros-v2.json)
Output: a .docx and a .md, same basename. The `v2` marker lives in the FILENAME
        and Doc title only - the rendered content never says "v2".

Cover-page deltas vs the template render (the ONLY body/cover differences, both
deliberate - a template is generic, a Client ROS is one firm's recording copy):
  - the firm name renders under the episode title
  - the recording date renders under `Prepared by Case Engine` (omitted while TBD)

Dependencies: python-docx.

Usage:
  python3 build-client-ros-v2-docx.py \
      --data client-ros-v2-data.json \
      --output "E7: Truck Accidents // Spaulding Injury Law // Client ROS v2 - GA - Savannah.docx"
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# CE brand (per Case Engine Branding folder colors.md)
CE_BLUE = RGBColor(0x35, 0x73, 0xFF)
CE_DARK = RGBColor(0x0F, 0x17, 0x2A)
CE_GRAY = RGBColor(0x5B, 0x66, 0x76)
FONT = "Roboto"

# Freestanding: nothing here may resolve to a path above SKILL_DIR. The logo is bundled.
SKILL_DIR = Path(__file__).resolve().parent.parent
DEFAULT_LOGO = SKILL_DIR / "assets" / "case-engine-logo.png"

PPR_AFTER_PBDR = (
    "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap", "w:overflowPunct",
    "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
    "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
    "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
    "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)


# ---------- the populate contract ----------

# The twelve approved v2 tokens. Mirror of
# pod-3A-ros-template-v2/references/placeholders.md - that file is canonical; if the two ever
# disagree, IT wins and this tuple is the bug.
V2_TOKENS = (
    "TOPIC", "CITY", "STATE", "PODCAST_NAME", "ATTORNEY_NAME", "ATTORNEY", "INTERVIEWER",
    "FIRM_NAME", "PHONE_NUMBER", "WEBSITE", "PODCAST_DOMAIN", "YEARS_PRACTICING",
)
# WEBSITE = the BUSINESS site, used in the conversion CTA (a case inquiry must land on the firm).
# PODCAST_DOMAIN = where episodes live, used in the subscribe line. Gabe directive 2026-08-21:
# these are two different things and must never collapse into one token.

# The 2 STATIC template constants (outro_note removed 2026-08-21 per Gabe), tokens unresolved. Mirror of
# pod-3A-ros-template-v2/references/statics.json v2.0.0 (canonical). The populated statics in the
# payload must equal these with tokens resolved - any other delta means a run regenerated
# boilerplate, which is the drift Editorial Guideline 8 exists to prevent.
TEMPLATE_STATICS = {
    "welcome": "Welcome back to **{{PODCAST_NAME}}** with **{{ATTORNEY_NAME}}**.",
    "welcome_first": "Welcome to **{{PODCAST_NAME}}** with **{{ATTORNEY_NAME}}**.",
    # Embedded-name variants (Gabe 2026-08-26): when the podcast name embeds the attorney's name
    # (e.g. "Car Accident Attorney w. Robert May"), "with **{{ATTORNEY_NAME}}**" would double the
    # name. A "w." in the podcast name is spoken, and rendered in the welcome, as "with".
    "welcome_embedded": "Welcome back to the **{{PODCAST_NAME}}** Podcast.",
    "welcome_embedded_first": "Welcome to the **{{PODCAST_NAME}}** Podcast.",
}

OUTRO_LINES = ("thanks", "signoff", "reach")

# Follow-ups: render under the S1 attribute block, never read on air. BOTH the interviewer
# AND the attorney read them, so they are written in the second person TO the attorney -
# never third-person instructions about him (Gabe 2026-08-24).
# The first bullet is HARD-CODED and byte-identical every episode (Gabe 2026-08-24); the
# remaining bullets are generated per topic and are where the case-study prompt lives.
FOLLOWUP_NOTE = "Not read on air. For the interviewer and the attorney both."
FOLLOWUP_STATIC = ("Follow up when the opportunity presents itself, not on a schedule. "
                   "Let the answer finish first.")

# The line under the S2 heading. Direction to the host, and it ships. Byte-identical to the
# template renderer's SF_NOTE.
SF_NOTE = ("These question answers are designed to be clipped in short form and answered in 60 "
           "to 90 seconds.")

TOKEN_RE = re.compile(r"\{\{[A-Z0-9_]+\}\}")
GUEST_RE = re.compile(r"\b(my guest|our guest|today'?s guest|joining us|thanks for coming on)\b",
                      re.I)


def resolve(text, values):
    """Resolve {{TOKEN}} -> value with str.replace, NEVER str.format.

    format() collapses the doubled braces and silently destroys every token it does not know.
    replace() touches only exact tokens and leaves everything else alone.
    """
    if not text:
        return text or ""
    for tok in V2_TOKENS:
        val = values.get(tok)
        if val is not None:
            text = text.replace("{{%s}}" % tok, str(val))
    return text


def walk_strings(obj, path=""):
    """Yield (path, string) for every string in a nested payload."""
    if isinstance(obj, str):
        yield path, obj
    elif isinstance(obj, dict):
        for k, v in obj.items():
            yield from walk_strings(v, f"{path}.{k}" if path else k)
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            yield from walk_strings(v, f"{path}[{i}]")


# ---------- inline markup (identical to the template renderer) ----------

PANDOC_COLOR_RE = re.compile(r"\{\.color=[^}]+\}")
PANDOC_CLASS_RE = re.compile(r"\{\.[a-z][a-zA-Z0-9_-]*(?:=[^}]*)?\}")
PANDOC_INLINE_RE = re.compile(r"\[([^\]]+)\]\{\.(underline|smallcaps|mark)\}")
UTAG_RE = re.compile(r"<u>(.*?)</u>", re.S)


def _is_punct(ch):
    return not ch.isalnum() and not ch.isspace()


def _flanking(text, start, end):
    prev = text[start - 1] if start > 0 else " "
    nxt = text[end] if end < len(text) else " "
    left = not nxt.isspace() and (not _is_punct(nxt) or prev.isspace() or _is_punct(prev))
    right = not prev.isspace() and (not _is_punct(prev) or nxt.isspace() or _is_punct(nxt))
    return left, right


def parse_inline(text, *, bold=False, italic=False, underline=False):
    """(chunk, bold, italic, underline) tuples. Depth-counting scan so nested bold (a populated
    value's own ** inside a fully bold Short-Form question) nests instead of closing early."""
    runs = []
    if not text:
        return runs
    buf = []
    bd = 0
    itd = 0

    def flush():
        if not buf:
            return
        plain = PANDOC_COLOR_RE.sub("", PANDOC_CLASS_RE.sub("", "".join(buf)))
        del buf[:]
        if plain:
            runs.append((plain, bold or bd > 0, italic or itd > 0, underline))

    pos = 0
    n = len(text)
    while pos < n:
        ch = text[pos]
        if ch == "[":
            m = PANDOC_INLINE_RE.match(text, pos)
            if m:
                flush()
                runs.extend(parse_inline(m.group(1), bold=bold or bd > 0,
                                         italic=italic or itd > 0, underline=True))
                pos = m.end()
                continue
        elif ch == "<":
            m = UTAG_RE.match(text, pos)
            if m:
                flush()
                runs.extend(parse_inline(m.group(1), bold=bold or bd > 0,
                                         italic=italic or itd > 0, underline=True))
                pos = m.end()
                continue
        elif ch == "*":
            width = 2 if pos + 1 < n and text[pos + 1] == "*" else 1
            left, right = _flanking(text, pos, pos + width)
            if width == 2 and bd > 0 and right:
                flush(); bd -= 1; pos += 2; continue
            if width == 2 and left:
                flush(); bd += 1; pos += 2; continue
            if width == 1 and itd > 0 and right:
                flush(); itd -= 1; pos += 1; continue
            if width == 1 and left:
                flush(); itd += 1; pos += 1; continue
        buf.append(ch)
        pos += 1
    flush()
    return runs


def strip_pandoc(text):
    if text is None:
        return ""
    text = PANDOC_INLINE_RE.sub(r"\1", text)
    text = PANDOC_COLOR_RE.sub("", text)
    return PANDOC_CLASS_RE.sub("", text)


# ---------- docx primitives (identical to the template renderer) ----------

def set_run_style(run, *, size=11, bold=False, italic=False, underline=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def add_para(doc, text="", *, align=None, space_before=0, space_after=6, style=None, **run_kw):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_run_style(p.add_run(text), **run_kw)
    return p


def add_rich(doc, text, *, base_size=11, base_bold=False, base_italic=False,
             base_color=CE_DARK, space_before=0, space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for chunk, bo, it, un in parse_inline(text or ""):
        set_run_style(p.add_run(chunk), size=base_size, bold=bo or base_bold,
                      italic=it or base_italic, underline=un, color=base_color)
    return p


def add_h1(doc, text):
    return add_rich(doc, text, base_size=20, base_bold=True, base_color=CE_BLUE,
                    space_before=18, space_after=8)


def add_h2(doc, text):
    return add_rich(doc, text, base_size=16, base_bold=True, base_color=CE_DARK,
                    space_before=14, space_after=6)


def add_note(doc, text):
    return add_rich(doc, text, base_size=11, base_italic=True, base_color=CE_GRAY, space_after=6)


def add_bullet(doc, text):
    return add_rich(doc, text, base_size=11, base_color=CE_DARK, space_after=3,
                    style="List Bullet")


def add_sf_location(doc, text):
    return add_rich(doc, text, base_size=16, base_bold=True, base_color=CE_DARK,
                    space_before=8, space_after=2)


def add_sf_question(doc, text):
    return add_rich(doc, text, base_size=11, base_bold=True, base_color=CE_DARK,
                    space_before=0, space_after=0)


def add_sf_bullet(doc, text):
    return add_rich(doc, text, base_size=11, base_color=CE_DARK, space_before=0, space_after=0,
                    style="List Bullet")


def add_rule(doc, *, space_before=12, space_after=12):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = False
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "5B6676")
    pbdr.append(bottom)
    p._element.get_or_add_pPr().insert_element_before(pbdr, *PPR_AFTER_PBDR)
    return p


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _hdr_ftr_text(container, text, *, align=WD_ALIGN_PARAGRAPH.LEFT, size=9):
    p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.text = ""
    p.alignment = align
    set_run_style(p.add_run(text), size=size, color=CE_GRAY)


# ---------- payload helpers ----------

def scope_label(data):
    parts = [data.get("scope") or "Topic Only"]
    if data.get("location"):
        parts.append(data["location"])
    return " - ".join(parts)


def sf_question(q, where):
    text = (q.get("q") or q.get("question") or "").strip()
    if not text:
        raise SystemExit(f"FAILED: {where} carries no question text.")
    return text.replace("**", "")


def sf_bullets(q):
    out = []
    for b in q.get("bullets") or []:
        if isinstance(b, str):
            if b.strip():
                out.append(b.strip())
            continue
        label = (b.get("label") or "").strip()
        detail = (b.get("detail") or "").strip()
        if label:
            out.append(f"[{label}]{{.underline}}: {detail}".rstrip(": "))
        elif detail:
            out.append(detail)
    return out


def outro_lines(data):
    outro = data.get("outro") or {}
    missing = [k for k in OUTRO_LINES if not outro.get(k)]
    if missing:
        raise SystemExit(f"FAILED: outro block missing generated line(s): {missing}.")
    return [outro[k] for k in OUTRO_LINES]


# ---------- populate gates ----------

def firm_values(data):
    """The twelve resolved values as a TOKEN -> value map, from the payload's firm block."""
    firm = data.get("firm") or {}
    mapping = {
        "TOPIC": data.get("topic_spoken") or data.get("topic"),
        "CITY": data.get("city"),
        "STATE": data.get("state"),
        "PODCAST_NAME": firm.get("podcast_name"),
        "ATTORNEY_NAME": firm.get("attorney_name"),
        "ATTORNEY": firm.get("attorney"),
        "INTERVIEWER": firm.get("interviewer"),
        "FIRM_NAME": firm.get("firm_name"),
        "PHONE_NUMBER": firm.get("phone_number"),
        "WEBSITE": firm.get("website"),
        "PODCAST_DOMAIN": firm.get("podcast_domain"),
        "YEARS_PRACTICING": firm.get("years_practicing"),
    }
    return mapping


def validate_populated(data):
    """The populate gates. Every one is hard - a failure writes no file.

    1. episode_format is v2-open-interview (this renderer never touches legacy).
    2. Zero leftover {{...}} anywhere in the payload - an unresolved token would be read aloud.
    3. The statics equal the template constants with tokens resolved - not regenerated.
    4. No appendix in the payload - the Source Question Bank is INTERNAL to the template.
    5. Attribute block: 10-12 bullets, zero question marks (AT-1 / AT-2, inherited).
    6. Eight to ten questions per location, 2-4 bullets each (Gabe directive 2026-08-21;
       was ten. The n-gram table ships 20/location: 15 render here, 5 to the Question Pool).
    7. Zero guest framing, zero em dashes (inherited).
    """
    if data.get("episode_format") != "v2-open-interview":
        raise SystemExit(
            "FAILED: episode_format is not 'v2-open-interview'. This renderer only builds v2 "
            "Client ROS docs. A legacy Client ROS belongs to pod-3B-client-ros."
        )

    leftovers = []
    for path, s in walk_strings({k: v for k, v in data.items() if k != "metadata"}):
        for m in TOKEN_RE.finditer(s):
            leftovers.append(f"{path}: {m.group(0)}")
    if leftovers:
        raise SystemExit(
            "FAILED: unresolved {{...}} token(s) in the populated payload - these would ship "
            "into the recording as literal markup:\n  " + "\n  ".join(leftovers[:20])
        )

    if "appendix_question_bank" in data and data["appendix_question_bank"]:
        raise SystemExit(
            "FAILED: payload carries appendix_question_bank. The Source Question Bank is "
            "INTERNAL to the ROS Template and never reaches the Client ROS. Strip it."
        )

    values = firm_values(data)
    missing = [t for t, v in values.items() if v in (None, "")]
    if missing:
        raise SystemExit(f"FAILED: firm block missing resolved value(s) for: {missing}")

    # The payload stores the ACTIVE welcome under `static.welcome` - the Episode 1 variant when
    # is_first_episode is true, the standard one otherwise. Both must equal the template
    # constant with tokens resolved. outro_note was removed from the STATIC set 2026-08-21.
    S = data.get("static") or {}
    first = bool(data.get("is_first_episode"))
    welcome_key = "welcome_first" if first else "welcome"
    embedded_key = "welcome_embedded_first" if first else "welcome_embedded"
    # Either constant is acceptable (Gabe 2026-08-26): the standard welcome, or the embedded-name
    # variant when the podcast name embeds the attorney's name. A "w." in the podcast name is
    # spoken, and rendered in the welcome, as "with", so the embedded variant is also accepted
    # with that spoken form applied.
    spoken = dict(values)
    if spoken.get("PODCAST_NAME"):
        spoken["PODCAST_NAME"] = re.sub(r"\bw\.(?=\s)", "with", spoken["PODCAST_NAME"])
    acceptable = {
        resolve(TEMPLATE_STATICS[welcome_key], values).strip(),
        resolve(TEMPLATE_STATICS[embedded_key], values).strip(),
        resolve(TEMPLATE_STATICS[embedded_key], spoken).strip(),
    }
    got = (S.get("welcome") or "").strip()
    if got not in acceptable:
        raise SystemExit(
            "FAILED: STATIC 'welcome' does not equal a template constant with tokens "
            f"resolved.\n  expected one of: {sorted(acceptable)}\n  got:             {got}\n"
            "A run regenerated boilerplate (Editorial Guideline 8). Populate is .replace() "
            "on the constant, never regeneration."
        )

    s1 = data["segment_1"]
    fups = s1.get("follow_ups") or []
    if not fups:
        raise SystemExit("FAILED: segment_1.follow_ups is empty; at least one topic-specific "
                         "follow-up is required (the case-study prompt).")
    if any("?" in f for f in fups):
        raise SystemExit("FAILED: a follow-up carries a question mark; these are interviewer "
                         "notes, not lines to read.")

    attrs = s1.get("attributes") or []
    if not 10 <= len(attrs) <= 12:
        raise SystemExit(f"FAILED: attribute block has {len(attrs)} bullet(s); 10-12 required (AT-2).")
    for ai, a in enumerate(attrs, start=1):
        if "?" in (a.get("name") or "") or "?" in (a.get("detail") or ""):
            raise SystemExit(f"FAILED: attribute bullet {ai} carries a question mark (AT-1).")

    for block in data["segment_2"]["locations"]:
        loc = block.get("location") or "(unnamed)"
        questions = block.get("questions") or []
        if not 8 <= len(questions) <= 10:
            raise SystemExit(f"FAILED: location '{loc}' has {len(questions)} questions; 8 to 10 required "
                         f"(Gabe 2026-08-24; was 15). The n-gram table still ships 20 per location - "
                         f"the best 8 to 10 render here, the rest go to the Question Pool.")
        for qi, q in enumerate(questions, start=1):
            n = len(sf_bullets(q))
            if not 2 <= n <= 4:
                raise SystemExit(f"FAILED: '{loc}' Q{qi} has {n} bullet(s); 2-4 required.")

    for path, s in walk_strings({k: v for k, v in data.items() if k != "metadata"}):
        if GUEST_RE.search(s):
            raise SystemExit(f"FAILED: guest framing at {path}: '{GUEST_RE.search(s).group(0)}'. "
                             "The attorney owns the show.")
        if "—" in s or "–" in s:
            raise SystemExit(f"FAILED: em/en dash at {path}. Plain hyphens only.")


# ---------- document ----------

def build_docx(data, logo):
    """The locked v2 shape, populated. Byte-identical structure to the template render minus the
    appendix, plus the two cover deltas (firm name line, recording date line)."""
    doc = Document()
    firm = data["firm"]
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        _hdr_ftr_text(section.header,
                      f"Case Engine  |  Run of Show  |  {firm['firm_name']}  |  "
                      f"{data['topic']}, {scope_label(data)}")
        _hdr_ftr_text(section.footer, "Case Engine  |  Confidential")

    S = data["static"]
    s1, s2 = data["segment_1"], data["segment_2"]
    cover = data.get("cover_page", {})

    # ---- cover page ----
    doc.add_paragraph()
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo:
        logo_p.add_run().add_picture(logo, width=Pt(180))
    doc.add_paragraph()

    add_para(doc, cover.get("title", "Run of Show"), align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=6, size=24, bold=True, color=CE_BLUE)
    add_para(doc, data["episode_title"], align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=18, bold=True, color=CE_DARK)
    # Cover delta 1: the firm. A Client ROS is one firm's recording copy, so the firm is on
    # the cover - the template, being generic, never carries it.
    add_para(doc, firm["firm_name"], align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=14, bold=True, color=CE_DARK)
    doc.add_paragraph()
    add_para(doc, f"{data['topic']}  |  {scope_label(data)}", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=14, color=CE_DARK)
    doc.add_paragraph()
    add_para(doc, cover.get("prepared_by", "Prepared by Case Engine"),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, size=11, color=CE_DARK)
    # Cover delta 2: the recording date, per firm, collected by this skill. Omitted while TBD -
    # never print "Recording: TBD" on a client-facing cover.
    rec = (firm.get("recording_date") or "").strip()
    if rec and rec.upper() != "TBD":
        add_para(doc, f"Recording: {rec}", align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=2, size=11, color=CE_GRAY)
    add_page_break(doc)

    # ---- S1: Long-Form ----
    add_h1(doc, "S1: Long-Form (15-30m)")
    add_h2(doc, "Introduction")
    add_rich(doc, S["welcome"])
    add_rich(doc, s1.get("setup") or s1.get("cold_open") or "")
    if s1.get("credential"):
        add_rich(doc, s1["credential"])
    add_rich(doc, s1["prompt"], base_bold=True)
    add_h2(doc, "ATTORNEY RESPONSE")
    # Bullets, never a numbered list (Gabe 2026-08-26).
    for a in s1["attributes"]:
        add_bullet(doc, f"**{a['name']}.** {a['detail']}")

    fups = s1.get("follow_ups") or []
    add_h2(doc, "Follow-ups")
    add_note(doc, FOLLOWUP_NOTE)
    add_bullet(doc, FOLLOWUP_STATIC)
    for f in fups:
        add_bullet(doc, f)

    add_h2(doc, "Outro")
    for line in outro_lines(data):
        add_rich(doc, line)

    # ---- S2: Short-Form (flows on behind a rule, never its own page) ----
    add_rule(doc)
    add_h1(doc, "S2: Short-Form (60-90s)").paragraph_format.page_break_before = False
    add_note(doc, SF_NOTE)
    for block in s2["locations"]:
        add_sf_location(doc, f"Location: {block['location']}")
        for qi, q in enumerate(block["questions"], start=1):
            add_sf_question(doc, f"**Q{qi}: {sf_question(q, block['location'])}**")
            for bullet in sf_bullets(q):
                add_sf_bullet(doc, bullet)

    # ---- NO appendix. The Source Question Bank is INTERNAL to the ROS Template. ----
    return doc


def build_markdown(data):
    S = data["static"]
    s1, s2 = data["segment_1"], data["segment_2"]
    firm = data["firm"]
    cover = data.get("cover_page", {})
    P = strip_pandoc
    out = []
    w = out.append

    w("# Run of Show")
    w("")
    w(f"**{data['episode_title']}**")
    w("")
    w(f"**{firm['firm_name']}**")
    w("")
    w(f"{data['topic']}  |  {scope_label(data)}")
    w("")
    w(cover.get("prepared_by", "Prepared by Case Engine"))
    rec = (firm.get("recording_date") or "").strip()
    if rec and rec.upper() != "TBD":
        w("")
        w(f"Recording: {rec}")
    w("")

    w("# S1: Long-Form (15-30m)")
    w("")
    w("## Introduction")
    w("")
    w(P(S["welcome"]))
    w("")
    w(P(s1.get("setup") or s1.get("cold_open") or ""))
    w("")
    if s1.get("credential"):
        w(P(s1["credential"])); w("")
    w(f"**{P(s1['prompt']).replace('**', '')}**")
    w("")
    w("## ATTORNEY RESPONSE")
    w("")
    # Bullets, never a numbered list (Gabe 2026-08-26).
    for a in s1["attributes"]:
        w(f"- **{a['name']}.** {P(a['detail'])}")
    w("")
    w("## Follow-ups")
    w("")
    w(f"*{FOLLOWUP_NOTE}*")
    w("")
    w(f"- {FOLLOWUP_STATIC}")
    for f in (s1.get("follow_ups") or []):
        w(f"- {P(f)}")
    w("")
    w("## Outro")
    w("")
    w("")
    for line in outro_lines(data):
        w(P(line)); w("")

    w("---")
    w("")
    w("# S2: Short-Form (60-90s)")
    w("")
    w(f"*{SF_NOTE}*")
    w("")
    for block in s2["locations"]:
        w(f"## Location: {block['location']}")
        w("")
        for qi, q in enumerate(block["questions"], start=1):
            w(f"**Q{qi}: {P(sf_question(q, block['location']))}**")
            w("")
            for bullet in sf_bullets(q):
                w(f"- {P(bullet)}")
            w("")
    return "\n".join(out)


# ---------- entry point ----------

def build(args):
    data = json.loads(Path(args.data).read_text())
    validate_populated(data)

    logo = args.logo or (str(DEFAULT_LOGO) if DEFAULT_LOGO.is_file() else None)
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    build_docx(data, logo).save(out)

    md_path = out.with_suffix(".md")
    md_path.write_text(build_markdown(data))

    locs = len(data["segment_2"]["locations"])
    print(f"DOCX: {out}")
    print(f"MD:   {md_path}")
    print(f"Data: {data['firm']['firm_name']} | {locs} location set(s) x 8-10 questions | "
          f"{len(V2_TOKENS)}/{len(V2_TOKENS)} tokens resolved | appendix stripped")
    if not logo:
        print(f"WARN: no --logo given and no bundled logo at {DEFAULT_LOGO}, cover page "
              "rendered without the CE logo")


def main():
    p = argparse.ArgumentParser(description="Build the CE-branded Client ROS v2 DOCX + .md")
    p.add_argument("--data", required=True, help="Path to client-ros-v2-data.json")
    p.add_argument("--logo", default=None,
                   help="CE logo path; defaults to bundled assets/case-engine-logo.png")
    p.add_argument("--output", required=True, help="Output .docx path (the .md is written alongside)")
    build(p.parse_args())


if __name__ == "__main__":
    sys.exit(main())
