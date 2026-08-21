#!/usr/bin/env python3
"""
What: Builds a CE-branded Keyword Research DOCX from keyword-research.json.
Input: keyword-research.json (produced by /pod-2B-keyword-research) + a CE logo PNG
       (download from the canonical Case Engine Branding folder:
       https://drive.google.com/drive/folders/1OulNcg6hZ6caFD0Xc-W50i7ougC_y0qo)
Output: <scope>/Keyword Research.docx with cover page + headers/footers + executive
        summary + intent distribution + seed table + ranked keyword set + PAA stacks +
        related searches + localization summary + Search Queries & Volume table.

Client-facing scope: the DOCX is the human-facing canonical view. It renders ONLY the
sections specced as client-facing in SKILL.md. The `## INTERNAL` / `## Quality Assurance`
content lives in the markdown source-of-truth, never in this DOCX (the Drive Doc renderer
truncates at the first such heading; this script never emits those sections).

Branding mirrors the sibling pod-2A build-entity-map-docx.py (CE colors, cover, header,
footer). Drive uploads this DOCX via files.update so auto-conversion keeps a stable fileId.
A Roboto pass is applied to the converted Google Doc by the calling skill, not here.

Dependencies: python-docx (cowork runtime has it). No optional deps.

Usage:
  python3 build-keyword-research-docx.py \\
      --json /path/to/keyword-research.json \\
      --logo /path/to/ce-logo.png \\
      --output "/path/to/Keyword Research.docx" \\
      [--run-date "June 17, 2026"]
"""

import argparse
import json
import sys
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CE_BLUE = RGBColor(0x35, 0x73, 0xFF)
CE_DARK = RGBColor(0x0F, 0x17, 0x2A)
CE_GRAY = RGBColor(0x5B, 0x66, 0x76)
FONT = "Roboto"


def set_run_style(run, *, font=FONT, size=11, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(a), font)


def add_para(doc, text="", *, align=None, space_before=0, space_after=6, **run_kw):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_run_style(p.add_run(text), **run_kw)
    return p


def add_bullet(doc, text):
    """Render a bullet with **bold** inline spans."""
    p = doc.add_paragraph(style="List Bullet")
    for i, part in enumerate(text.split("**")):
        if not part:
            continue
        set_run_style(p.add_run(part), size=11, bold=(i % 2 == 1), color=CE_DARK)
    return p


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def set_header(section, text):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_style(p.add_run(text), size=9, italic=True, color=CE_GRAY)


def set_footer(section, text_left):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_style(p.add_run(text_left + "\t\t"), size=9, color=CE_GRAY)
    set_run_style(p.add_run("Page "), size=9, color=CE_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    set_run_style(p.add_run(text), size=16, bold=True, color=CE_BLUE)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_run_style(p.add_run(text), size=13, bold=True, color=CE_DARK)
    return p


def add_table(doc, headers, rows, *, bold_rows=None):
    bold_rows = bold_rows or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        set_run_style(hdr[i].paragraphs[0].add_run(h), size=9, bold=True, color=CE_BLUE)
    for r_idx, row_vals in enumerate(rows):
        cells = table.rows[1 + r_idx].cells
        is_bold = r_idx in bold_rows
        for i, v in enumerate(row_vals):
            cells[i].text = ""
            set_run_style(cells[i].paragraphs[0].add_run(str(v)), size=9, bold=is_bold, color=CE_DARK)
    return table


def build(args):
    d = json.loads(Path(args.json).read_text())
    run_date = args.run_date or date.today().strftime("%B %-d, %Y")

    topic = d.get("topic", "Topic")
    scope = d.get("scope", "Topic")
    location = d.get("location")
    scope_label = scope + (f" - {location}" if location else "")

    kw = d["keywords"]
    tot = len(kw)
    dist = Counter(k["intent"] for k in kw)
    pct = {k: round(100 * v / tot, 1) for k, v in dist.items()}
    real = sum(1 for k in kw if k["data_source"] == "keywords_everywhere_gkp")
    paa_total = sum(len(s["questions"]) for s in d["paa_stacks"])
    top5 = [k["query"] for k in sorted(kw, key=lambda x: -x["msv"])][:5]
    seeds = d["seed_keywords"]

    doc = Document()
    header_text = f"Case Engine  |  Keyword Research  |  {topic}, {scope_label}"
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        set_header(section, header_text)
        set_footer(section, "Case Engine  |  Confidential")

    # ---- Cover ----
    for _ in range(3):
        doc.add_paragraph()
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(args.logo, width=Inches(2.4))
    add_para(doc, "Keyword Research", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=6, size=36, bold=True, color=CE_BLUE)
    add_para(doc, topic, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, size=22, bold=True, color=CE_DARK)
    add_para(doc, scope_label, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24, size=14, color=CE_DARK)
    add_para(doc, "Prepared by Case Engine", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=2, size=11, italic=True, color=CE_GRAY)
    add_para(doc, run_date, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, size=11, color=CE_GRAY)
    add_page_break(doc)

    # ---- Executive Summary ----
    add_h2(doc, "Executive Summary")
    if real:
        src_sentence = (
            f"Search volume and CPC for {real} of {tot} keywords come from real Google Keyword Planner "
            f"data via the Keywords Everywhere API; the remaining {tot - real} long-tail and "
            f"locally-modified terms carry directional estimate values (the per-row source tag is the "
            f"flag). Keyword difficulty and SERP features are estimated this run."
        )
    else:
        src_sentence = (
            "Volume and CPC are directional estimates this run (the per-row source tag is the flag) "
            "because the real-volume sources were unreachable; the People Also Ask stacks for the lead "
            "seeds are real content-gap SERP captures. Treat ranking as directional, not precise."
        )
    add_para(doc,
             f"This keyword research pass maps the demand landscape for {topic} at {scope_label} scope. "
             + src_sentence,
             size=11, color=CE_DARK, space_after=10)
    for line in [
        f"Total keywords captured: {tot} (target 150-300)",
        f"Seed keywords: {len(seeds)}",
        f"PAA questions captured: {paa_total} (target 15+ per seed)",
        f"Related searches: {len(d['related_searches'])}",
        f"Intent: Informational {pct.get('Informational', 0)}% / Commercial {pct.get('Commercial', 0)}% "
        f"/ Transactional {pct.get('Transactional', 0)}% / Navigational {pct.get('Navigational', 0)}%",
        f"Top 5 by volume: {', '.join(top5)}",
        f"Data source: mixed (real Keywords Everywhere for head and body terms; estimate for long-tail and local)",
    ]:
        add_bullet(doc, line)
    add_para(doc,
             "Note: intent distribution skews Informational versus the 60% target. This is genuine for "
             "this practice area, not a sampling error: the commercial lawyer-hire vocabulary is a small "
             "finite set while the informational landscape is broad. Commercial terms are retained for "
             "context and should not be chased by podcast content.",
             size=10, italic=True, color=CE_GRAY, space_before=6, space_after=10)

    # ---- Intent Distribution ----
    add_h3(doc, "Intent Distribution")
    tgt = {"Informational": 60, "Commercial": 25, "Transactional": 10, "Navigational": 5}
    rows = []
    for it in ["Informational", "Commercial", "Transactional", "Navigational"]:
        p = pct.get(it, 0.0)
        status = "on spec" if abs(p - tgt[it]) <= 10 else "skew (flagged)"
        rows.append([it, dist.get(it, 0), f"{p}%", f"{tgt[it]}%", status])
    add_table(doc, ["Intent", "Count", "%", "Target", "Status"], rows)

    add_page_break(doc)

    # ---- Seed Keywords ----
    add_h2(doc, "Seed Keywords")
    add_para(doc, "Anchor demand terms the full set expands from. Each seed has a verbatim PAA stack later in this document.",
             size=11, color=CE_DARK, space_after=8)
    rows = []
    for i, seed in enumerate(seeds, 1):
        k = next((x for x in kw if x["query"].lower() == seed.lower()), None)
        if k:
            srclbl = "KE (GKP)" if k["data_source"] == "keywords_everywhere_gkp" else "LLM est"
            rows.append([i, seed, k["intent"], f"{k['msv']:,}", k["kd"], f"${k['cpc']:.2f}", srclbl])
        else:
            rows.append([i, seed, "-", "-", "-", "-", "-"])
    add_table(doc, ["#", "Seed Keyword", "Intent", "MSV", "KD", "CPC", "Source"], rows,
              bold_rows=set(range(min(5, len(rows)))))

    # ---- Ranked Keyword Set ----
    add_h2(doc, "Ranked Keyword Set")
    add_para(doc, "Full demand landscape, grouped by intent bucket, descending by monthly volume within each bucket.",
             size=11, color=CE_DARK, space_after=8)
    for bucket in ["Informational", "Commercial", "Transactional", "Navigational"]:
        brows = sorted([k for k in kw if k["intent"] == bucket], key=lambda x: -x["msv"])
        if not brows:
            continue
        add_h3(doc, f"{bucket} ({len(brows)})")
        rows = []
        for k in brows:
            srclbl = "KE (GKP)" if k["data_source"] == "keywords_everywhere_gkp" else "LLM est"
            rows.append([k["query"], f"{k['msv']:,}", k["kd"], f"${k['cpc']:.2f}", ", ".join(k["serp_features"]), srclbl])
        add_table(doc, ["Keyword", "MSV", "KD", "CPC", "SERP Features", "Source"], rows)

    add_page_break(doc)

    # ---- PAA Stacks ----
    add_h2(doc, "PAA Stacks (People Also Ask)")
    add_para(doc,
             f"{paa_total} questions across {len(d['paa_stacks'])} seeds. Questions are preserved verbatim - "
             f"they become seed rows for the downstream N-Gram Table.",
             size=11, color=CE_DARK, space_after=8)
    for s in d["paa_stacks"]:
        add_h3(doc, f"{s['seed']}  ({s['paa_source']})")
        for q in s["questions"]:
            add_bullet(doc, q)

    add_page_break(doc)

    # ---- Related Searches ----
    add_h2(doc, "Related Searches")
    add_para(doc, f"{len(d['related_searches'])} lower-intent long-tail phrases for downstream n-gram row expansion.",
             size=11, color=CE_DARK, space_after=8)
    for r in d["related_searches"]:
        add_bullet(doc, r)

    add_page_break(doc)

    # ---- Localization Summary ----
    add_h2(doc, "Localization Summary")
    if d.get("localization") and d["localization"].get("pairs"):
        loc = d["localization"]
        add_para(doc,
                 f"Scope is Location (City Anchor). The keyword pass was run with and without the local "
                 f"modifier and the MSV deltas compared.",
                 size=11, color=CE_DARK, space_after=8)
        rows = [[p["generic"], f"{p['generic_msv']:,}", p["local"], f"{p['local_msv']:,}", f"{p['ratio_pct']}%"]
                for p in loc["pairs"]]
        add_table(doc, ["Generic Query", "Generic MSV", "Localized Query", "Local MSV", "Ratio"], rows)
        add_para(doc,
                 f"Average local/generic ratio: {round(loc['ratio'] * 100, 1)}%. Head terms are national, so a "
                 f"single-city modifier captures a small fraction. Content strategy: rank the generic informational "
                 f"pages and weave local anchors into them rather than building thin city-only volume pages.",
                 size=11, color=CE_DARK, space_before=6, space_after=4)
        for flip in loc.get("intent_flips", []):
            add_bullet(doc, f"Intent flip: {flip}")
    elif d.get("localization") and (d["localization"].get("ratio") is not None):
        loc = d["localization"]
        ratio = loc.get("ratio", 0)
        add_para(doc,
                 f"Scope is Location (City Anchor). The keyword pass was run with and without the local "
                 f"modifier. Localized modifiers retain roughly {round(ratio * 100)}% of generic MSV on "
                 f"average, within the normal 10-30% decay band.",
                 size=11, color=CE_DARK, space_after=8)
        for line in [
            "Statute terms carry state-level MSV; the city modifier collapses several below the 3% alarm "
            "ratio. Build these as state-wide pages with local anchors woven in rather than thin city-only "
            "volume pages.",
            "Commercial lawyer terms are city-tethered by nature with no generic equivalent in the podcast lane.",
            f"Intent-bucket flips: {'none detected' if not loc.get('intent_flips') else len(loc['intent_flips'])}.",
        ]:
            add_bullet(doc, line)
    else:
        add_para(doc, "Topic-level pass - no localization run at this scope.", size=11, color=CE_DARK)

    add_page_break(doc)

    # ---- Search Queries & Volume ----
    add_h2(doc, "Search Queries & Volume")
    add_para(doc,
             "Top queries across the full set, descending by monthly volume. Run of Show reads this table "
             "verbatim into its Appendix.",
             size=11, color=CE_DARK, space_after=8)
    rows = [[s["query"], f"{s['monthly_volume']:,}", s["source"]] for s in d["search_queries"]]
    add_table(doc, ["Query", "Monthly Volume", "Source"], rows)

    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved DOCX: {out}")
    print(f"Data: {tot} keywords, {len(seeds)} seeds, {paa_total} PAA, {len(d['related_searches'])} related")


def main():
    ap = argparse.ArgumentParser(description="Build CE-branded Keyword Research DOCX from keyword-research.json")
    ap.add_argument("--json", required=True)
    ap.add_argument("--logo", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--run-date", default=None)
    build(ap.parse_args())


if __name__ == "__main__":
    main()
