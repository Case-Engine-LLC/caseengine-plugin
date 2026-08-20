#!/usr/bin/env python3
"""
What: Builds a CE-branded ROS Template DOCX (and paired .md sibling) from a
      ros-template-data.json payload that matches _references/schemas/ros-template.json.
Input: ros-template-data.json (filename per the canonical Podcast Drive doc:
       https://docs.google.com/document/d/1YhybGpp9DIqmV56P6OOHIQe7A6RxvjQGHLHKcUM0JmU)
       + a CE logo image (download from the canonical Case Engine Branding folder:
       https://drive.google.com/drive/folders/1OulNcg6hZ6caFD0Xc-W50i7ougC_y0qo)
Output: ROS Template.docx with cover page + headers/footers + Producer Notes +
        Introduction + Segments S1..N + Closing/CTA, paired alongside a clean
        ROS Template.md (raw markdown - same content shape, pandoc artifacts
        stripped to plain text).
Re-run: Safe - overwrites the output DOCX + .md. Drive's auto-conversion preserves
        the same fileId when the calling skill uses files.update for re-uploads
        (per the Push to Drive doc).

Convention sync: section order, branding spec, and the canonical filename live in the
Podcast Drive doc + the Case Engine Branding folder. If those change, update this script
to match. The canonical sources are the single source of truth.

Tokenized layer: this script renders the GENERIC reusable script. Every firm-specific
value is preserved as one of the 12 approved {{PLACEHOLDERS}} (see SKILL.md > Best
Practices > Placeholder taxonomy). Tokens pass through verbatim. /client-ros resolves
tokens downstream against a specific firm's profile.

Excluded: the legacy "Internal Setup / Complete and delete this section before sharing"
checklist is intentionally NOT rendered. The SOP's Step 0 state check covers that work
with no in-document residue.

Pandoc handling: input data may contain `[text]{.underline}`, `{.smallcaps}`, `{.mark}`,
`{.color=...}` markers. Underlines translate to native DOCX underline runs. In the .md
output, all pandoc inline markers are stripped to plain text so the markdown reads
cleanly without artifacts.

Dependencies: python-docx (cowork runtime has it). No optional deps.

Usage:
  python3 build-ros-template-docx.py \\
      --data "/path/to/01 Strategy/ros-template-data.json" \\
      --logo /path/to/ce-logo.png \\
      --practice-area "Car Accidents" \\
      --episode-topic "How to File a Car Accident Claim" \\
      --scope "Location" \\
      --location "GA - Savannah" \\
      --output "/path/to/{Firm} Podcast/Episodes/E2 - How to File a Car Accident Claim - GA Savannah/01 Strategy/ROS Template - E2 - How to File a Car Accident Claim - GA Savannah.docx" \\
      [--run-date "April 27, 2026"]    (defaults to today)

Notes on scope + location + destination:
  - --scope must be one of: "Topic Only", "Location", "Extension"
  - --location is required when --scope is "Location" or "Extension"
  - At "Topic Only" scope --location is omitted; cover + header drop the location line.
  - DESTINATION: at Location/Extension scope, write into the firm episode folder's
    "01 Strategy/" directory ALONGSIDE the Client ROS + Client Guide
    ("{Firm} Podcast/Episodes/E{N} - {Episode} - {Location}/01 Strategy/"), filename
    "ROS Template - E{N} - {Episode Short Title} - {Location}.docx" (append " (Extension)"
    for an extension cell). At Topic Only scope, write to the master templates tree:
    "templates [master]/AEO/Podcast/Episode Templates/{Topic}/{Episode}/Topic Only/06 ROS Template/ROS Template.docx".
  - The DOCX is the human-facing artifact: upload it to Drive as
    mimeType "application/vnd.google-apps.document" so Drive auto-converts to a clean
    branded Google Doc with REAL underlined entity runs (translated from pandoc
    [entity]{.underline} by this script). The paired .md is uploaded as text/markdown
    (raw source, no conversion). NEVER re-upload the .md with convert=true to make a
    second Doc - that leaks "[entity]{.underline}" as visible text and has no cover page.
  - Font is Roboto (see FONT constant; overrides Calibri per Gabe 2026-05-12).
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# CE brand (per Case Engine Branding folder colors.md)
CE_BLUE = RGBColor(0x35, 0x73, 0xFF)
CE_DARK = RGBColor(0x0F, 0x17, 0x2A)
CE_GRAY = RGBColor(0x5B, 0x66, 0x76)
FONT = "Roboto"  # CE branded deliverable font (overrides Calibri per Gabe 2026-05-12; if the Branding folder spec still says Calibri, Roboto wins)


# ---------- pandoc inline marker handling ----------

# Matches `[inner text]{.underline}` (and similar single-class markers).
# We capture the inner text + the class name so callers can decide what to do
# (translate to a styled DOCX run, or strip to plain text in the .md).
PANDOC_INLINE_RE = re.compile(r"\[([^\]]+)\]\{\.(underline|smallcaps|mark)\}")
# Matches `{.color=...}` style trailing markers attached to text - we strip these wholesale.
PANDOC_COLOR_RE = re.compile(r"\{\.color=[^}]+\}")
# Matches any other lingering `{.cls}` class-only markers - strip wholesale.
PANDOC_CLASS_RE = re.compile(r"\{\.[a-z][a-zA-Z0-9_-]*(?:=[^}]*)?\}")


def parse_pandoc_segments(text):
    """Split a string into a list of (text, class_or_None) tuples for DOCX rendering.

    `[Memorial Health]{.underline} is the Level I trauma center.` becomes:
      [("Memorial Health", "underline"), (" is the Level I trauma center.", None)]
    Trailing class-only markers like `{.color=red}` are stripped (treated as None-class plain text).
    """
    if text is None:
        return []
    segments = []
    pos = 0
    for m in PANDOC_INLINE_RE.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], None))
        segments.append((m.group(1), m.group(2)))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], None))
    # Wholesale-strip any color/class markers that survive in plain segments
    cleaned = []
    for seg_text, cls in segments:
        if cls is None:
            seg_text = PANDOC_COLOR_RE.sub("", seg_text)
            seg_text = PANDOC_CLASS_RE.sub("", seg_text)
        cleaned.append((seg_text, cls))
    return cleaned


def strip_pandoc(text):
    """Strip all pandoc inline markers and return plain text.

    Used when emitting the .md sibling - markdown should read cleanly without
    pandoc-specific artifacts. Tokens like {{FIRM_NAME}} pass through untouched.
    """
    if text is None:
        return ""
    # Replace `[inner]{.cls}` with just `inner`
    text = PANDOC_INLINE_RE.sub(r"\1", text)
    # Strip color and class-only markers
    text = PANDOC_COLOR_RE.sub("", text)
    text = PANDOC_CLASS_RE.sub("", text)
    return text


# ---------- unified inline parser (handles nested **bold** / *italic* / [text]{.underline}) ----------

# A combined token regex. Order matters: bold (**) before italic (*); pandoc spans;
# legacy <u>. Captures so we can recurse into bold/italic inner content for nested
# pandoc spans (e.g. `**[Empire Law]{.underline}**` -> bold + underlined run).
_INLINE_TOKEN_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*)"           # **bold** (non-greedy)
    r"|(\*(?P<ital>[^*]+?)\*)"           # *italic*
    r"|(\[(?P<pinner>[^\]]+)\]\{\.(?P<pcls>underline|smallcaps|mark)\})"   # [text]{.underline}
    r"|(<u>(?P<uinner>.*?)</u>)"         # <u>legacy</u>
)


def parse_inline(text, *, bold=False, italic=False, underline=False, smallcaps=False, mark=False):
    """Yield (chunk, bold, italic, underline, smallcaps, mark) tuples.

    Recursively descends into **bold** / *italic* spans so nested pandoc markers
    (e.g. the firm/attorney/podcast name wrapped `**[Name]{.underline}**`) render as
    a single run that is BOTH bold AND underlined - never leaking the literal
    `**` markers or the `[...]{.underline}` text.
    """
    if not text:
        return
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        start, end = m.span()
        if start > pos:
            # plain text between tokens - still strip any stray pandoc color/class markers
            plain = PANDOC_COLOR_RE.sub("", PANDOC_CLASS_RE.sub("", text[pos:start]))
            if plain:
                yield (plain, bold, italic, underline, smallcaps, mark)
        if m.group("bold") is not None:
            yield from parse_inline(m.group("bold"), bold=True, italic=italic,
                                    underline=underline, smallcaps=smallcaps, mark=mark)
        elif m.group("ital") is not None:
            yield from parse_inline(m.group("ital"), bold=bold, italic=True,
                                    underline=underline, smallcaps=smallcaps, mark=mark)
        elif m.group("pinner") is not None:
            cls = m.group("pcls")
            yield from parse_inline(m.group("pinner"), bold=bold, italic=italic,
                                    underline=underline or (cls == "underline"),
                                    smallcaps=smallcaps or (cls == "smallcaps"),
                                    mark=mark or (cls == "mark"))
        elif m.group("uinner") is not None:
            yield from parse_inline(m.group("uinner"), bold=bold, italic=italic,
                                    underline=True, smallcaps=smallcaps, mark=mark)
        pos = end
    if pos < len(text):
        plain = PANDOC_COLOR_RE.sub("", PANDOC_CLASS_RE.sub("", text[pos:]))
        if plain:
            yield (plain, bold, italic, underline, smallcaps, mark)


def _apply_smallcaps_mark(run, smallcaps, mark):
    if smallcaps:
        rPr = run._element.get_or_add_rPr()
        el = OxmlElement("w:smallCaps"); el.set(qn("w:val"), "1"); rPr.append(el)
    if mark:
        rPr = run._element.get_or_add_rPr()
        hl = OxmlElement("w:highlight"); hl.set(qn("w:val"), "yellow"); rPr.append(hl)


# ---------- DOCX run/paragraph helpers ----------

def set_run_style(run, *, font=FONT, size=11, bold=False, italic=False, underline=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)


def add_para(doc, text="", *, align=None, space_before=0, space_after=6, **run_kw):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_style(run, **run_kw)
    return p


def add_pandoc_para(doc, text, *, base_color=CE_DARK, base_size=11, base_bold=False,
                    base_italic=False, space_before=0, space_after=6, align=None):
    """Render `text` as a paragraph, translating ALL inline markers to native runs.

    Handles nested markdown bold/italic AND pandoc spans: `**[name]{.underline}**`
    becomes a single run that is bold + underlined; bare `**word**` becomes bold;
    `[name]{.underline}` becomes underlined. The literal `**`, `*`, `[...]{.underline}`,
    and `<u>` markers never leak into the rendered output. Tokens like {{FIRM_NAME}}
    pass through as plain text in their containing run.
    """
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for chunk, b, i, u, sc, mk in parse_inline(text, bold=base_bold, italic=base_italic):
        if not chunk:
            continue
        run = p.add_run(chunk)
        set_run_style(run, size=base_size, bold=b, italic=i, underline=u, color=base_color)
        _apply_smallcaps_mark(run, sc, mk)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_header(section, text):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_style(run, size=9, italic=True, color=CE_GRAY)


def set_footer(section, text_left):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text_left + "\t\t")
    set_run_style(run, size=9, color=CE_GRAY)
    run2 = p.add_run("Page ")
    set_run_style(run2, size=9, color=CE_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_style(run, size=16, bold=True, color=CE_BLUE)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_style(run, size=13, bold=True, color=CE_DARK)
    return p


def add_speaker_tag(doc, tag):
    """Render a `*[Interviewer]*` / `*[Attorney]*` italic speaker line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[{tag}]")
    set_run_style(run, size=11, italic=True, color=CE_DARK)
    return p


def add_question_text(doc, q_text):
    """Render the bolded question line under the H3. The whole line is bold; any
    pandoc `[entity]{.underline}` inside also gets the underline run property."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for chunk, b, i, u, sc, mk in parse_inline(q_text, bold=True):
        if not chunk:
            continue
        run = p.add_run(chunk)
        set_run_style(run, size=11, bold=True, italic=i, underline=u, color=CE_DARK)
        _apply_smallcaps_mark(run, sc, mk)
    return p


def add_bullet(doc, text, *, base_bold=False):
    """Render a bullet. Handles `**Label:**` prefixes and nested `[entity]{.underline}`
    spans (including `**[entity]{.underline}**`) via the unified inline parser - no
    literal `**` / `[...]{.underline}` markers leak."""
    p = doc.add_paragraph(style="List Bullet")
    for chunk, b, i, u, sc, mk in parse_inline(text, bold=base_bold):
        if not chunk:
            continue
        run = p.add_run(chunk)
        set_run_style(run, size=11, bold=b, italic=i, underline=u, color=CE_DARK)
        _apply_smallcaps_mark(run, sc, mk)
    return p


# ---------- DOCX build ----------

def build_docx(data, args, run_date):
    practice_area = args.practice_area
    episode_topic = args.episode_topic
    scope = args.scope
    location = args.location

    scope_label = f"{scope}" + (f" - {location}" if location else "")
    header_text = f"Case Engine  |  Run of Show  |  {practice_area}, {scope_label}"

    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        set_header(section, header_text)
        set_footer(section, "Case Engine  |  Confidential")

    # ---- Cover page ----
    for _ in range(3):
        doc.add_paragraph()

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(args.logo, width=Inches(2.4))

    add_para(doc, "Run of Show", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=6, size=36, bold=True, color=CE_BLUE)
    add_para(doc, episode_topic, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=22, bold=True, color=CE_DARK)
    add_para(doc, scope_label, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=24, size=14, color=CE_DARK)
    add_para(doc, "Prepared by Case Engine", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=2, size=11, italic=True, color=CE_GRAY)
    add_para(doc, run_date, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=0, size=11, color=CE_GRAY)

    add_page_break(doc)

    # ---- Producer Notes ----
    producer = data.get("producer_notes") or {}
    if producer:
        add_h2(doc, "Producer Notes")
        if producer.get("jurisdiction_summary"):
            add_pandoc_para(doc, producer["jurisdiction_summary"], space_after=8)
        if producer.get("attorney_bio_hooks"):
            add_pandoc_para(doc, producer["attorney_bio_hooks"], space_after=8)
        if producer.get("recording_notes"):
            add_pandoc_para(doc, producer["recording_notes"], base_italic=True,
                            base_color=CE_GRAY, space_after=8)

    # ---- Introduction ----
    intro = data.get("intro_block") or {}
    intro_dur = intro.get("duration_min", 2)
    add_h2(doc, f"Introduction (~{int(round(intro_dur))} minutes)")
    if intro.get("host_lines") or intro.get("topic_framing"):
        add_speaker_tag(doc, "Interviewer")
    if intro.get("host_lines"):
        for line in intro["host_lines"]:
            add_pandoc_para(doc, line, space_after=6)
    if intro.get("topic_framing"):
        add_pandoc_para(doc, intro["topic_framing"], space_after=6)

    # ---- Segments ----
    segments = data.get("segments", [])
    for s_idx, seg in enumerate(segments, start=1):
        s_name = seg.get("name", f"Segment {s_idx}")
        s_dur = seg.get("duration_min", 0)
        add_h2(doc, f"S{s_idx}: {s_name} (~{s_dur} minutes)")

        if seg.get("intro_prompt"):
            add_speaker_tag(doc, "Interviewer")
            add_pandoc_para(doc, seg["intro_prompt"], space_after=6)

        for q in seg.get("questions", []):
            q_id = q.get("q_id", "Q?")
            q_text = q.get("q_text", "")
            q_dur = q.get("duration_min", 0)
            add_h3(doc, f"{q_id}: {q_text} ({q_dur} minutes)")

            if q.get("setup"):
                add_speaker_tag(doc, "Interviewer")
                add_pandoc_para(doc, q["setup"], space_after=4)

            add_question_text(doc, q_text)

            add_speaker_tag(doc, "Attorney")
            for bullet in q.get("attorney_bullets", []):
                add_bullet(doc, bullet)

        if seg.get("transition_out"):
            add_speaker_tag(doc, "Interviewer")
            add_pandoc_para(doc, seg["transition_out"], space_after=6)

    # ---- Outro ----
    outro = data.get("outro_block") or {}
    if outro:
        outro_dur = outro.get("duration_min", 2)
        add_h2(doc, f"Closing and Call to Action (~{int(round(outro_dur))} minutes)")
        if outro.get("final_takeaway_prompt"):
            add_speaker_tag(doc, "Interviewer")
            add_pandoc_para(doc, outro["final_takeaway_prompt"], space_after=6)
        if outro.get("cta_lines"):
            for line in outro["cta_lines"]:
                add_speaker_tag(doc, "Interviewer")
                add_pandoc_para(doc, line, space_after=6)

    # ---- End marker ----
    add_para(doc, "End of Run of Show", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=18, space_after=6, size=11, italic=True, color=CE_GRAY)

    return doc


# ---------- Markdown sibling ----------

def build_markdown(data, args, run_date):
    """Build the raw .md deliverable - same content shape as the DOCX, plain markdown.

    Pandoc inline markers are stripped (no `[text]{.underline}` artifacts in the .md).
    Placeholder tokens pass through verbatim.
    """
    practice_area = args.practice_area
    episode_topic = args.episode_topic
    scope = args.scope
    location = args.location
    scope_label = f"{scope}" + (f" - {location}" if location else "")

    lines = []
    lines.append(f"# Run of Show: {strip_pandoc(episode_topic)}")
    lines.append("")
    lines.append(f"**Practice Area:** {strip_pandoc(practice_area)}")
    lines.append("")
    lines.append(f"**Scope:** {scope_label}")
    lines.append("")
    lines.append(f"_Prepared by Case Engine - {run_date}_")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Producer Notes
    producer = data.get("producer_notes") or {}
    if producer:
        lines.append("## Producer Notes")
        lines.append("")
        if producer.get("jurisdiction_summary"):
            lines.append(strip_pandoc(producer["jurisdiction_summary"]))
            lines.append("")
        if producer.get("attorney_bio_hooks"):
            lines.append(strip_pandoc(producer["attorney_bio_hooks"]))
            lines.append("")
        if producer.get("recording_notes"):
            lines.append(f"_{strip_pandoc(producer['recording_notes'])}_")
            lines.append("")
        lines.append("---")
        lines.append("")

    # Introduction
    intro = data.get("intro_block") or {}
    intro_dur = intro.get("duration_min", 2)
    lines.append(f"## Introduction (~{int(round(intro_dur))} minutes)")
    lines.append("")
    if intro.get("host_lines") or intro.get("topic_framing"):
        lines.append("*[Interviewer]*")
        lines.append("")
    if intro.get("host_lines"):
        for line in intro["host_lines"]:
            lines.append(strip_pandoc(line))
            lines.append("")
    if intro.get("topic_framing"):
        lines.append(strip_pandoc(intro["topic_framing"]))
        lines.append("")
    lines.append("---")
    lines.append("")

    # Segments
    for s_idx, seg in enumerate(data.get("segments", []), start=1):
        s_name = seg.get("name", f"Segment {s_idx}")
        s_dur = seg.get("duration_min", 0)
        lines.append(f"## S{s_idx}: {strip_pandoc(s_name)} (~{s_dur} minutes)")
        lines.append("")

        if seg.get("intro_prompt"):
            lines.append("*[Interviewer]*")
            lines.append("")
            lines.append(strip_pandoc(seg["intro_prompt"]))
            lines.append("")

        for q in seg.get("questions", []):
            q_id = q.get("q_id", "Q?")
            q_text = q.get("q_text", "")
            q_dur = q.get("duration_min", 0)
            lines.append(f"### {q_id}: {strip_pandoc(q_text)} ({q_dur} minutes)")
            lines.append("")

            if q.get("setup"):
                lines.append("*[Interviewer]*")
                lines.append("")
                lines.append(strip_pandoc(q["setup"]))
                lines.append("")

            # Bolded question text
            lines.append(f"**{strip_pandoc(q_text)}**")
            lines.append("")

            lines.append("*[Attorney]*")
            lines.append("")
            for bullet in q.get("attorney_bullets", []):
                lines.append(f"- {strip_pandoc(bullet)}")
            lines.append("")

        if seg.get("transition_out"):
            lines.append("*[Interviewer]*")
            lines.append("")
            lines.append(strip_pandoc(seg["transition_out"]))
            lines.append("")
        lines.append("---")
        lines.append("")

    # Outro
    outro = data.get("outro_block") or {}
    if outro:
        outro_dur = outro.get("duration_min", 2)
        lines.append(f"## Closing and Call to Action (~{int(round(outro_dur))} minutes)")
        lines.append("")
        if outro.get("final_takeaway_prompt"):
            lines.append("*[Interviewer]*")
            lines.append("")
            lines.append(strip_pandoc(outro["final_takeaway_prompt"]))
            lines.append("")
        if outro.get("cta_lines"):
            for line in outro["cta_lines"]:
                lines.append("*[Interviewer]*")
                lines.append("")
                lines.append(strip_pandoc(line))
                lines.append("")
        lines.append("---")
        lines.append("")

    lines.append("_End of Run of Show_")
    lines.append("")

    return "\n".join(lines)


# ---------- Orchestration ----------

def build(args):
    data_path = Path(args.data)
    if not data_path.exists():
        print(f"ERROR: {data_path} not found", file=sys.stderr)
        sys.exit(1)
    data = json.loads(data_path.read_text())

    run_date = args.run_date or date.today().strftime("%B %d, %Y")

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = build_docx(data, args, run_date)
    doc.save(str(output_path))
    print(f"Saved DOCX: {output_path}")

    md_path = output_path.with_suffix(".md")
    md_text = build_markdown(data, args, run_date)
    md_path.write_text(md_text)
    print(f"Saved MD:   {md_path}")

    seg_count = len(data.get("segments", []))
    q_count = sum(len(s.get("questions", [])) for s in data.get("segments", []))
    print(f"Data: {seg_count} segments, {q_count} questions")


def main():
    parser = argparse.ArgumentParser(
        description="Build CE-branded ROS Template DOCX + paired .md from ros-template-data.json"
    )
    parser.add_argument("--data", required=True,
                        help="Path to ros-template-data.json (matches _references/schemas/ros-template.json)")
    parser.add_argument("--logo", required=True,
                        help="Path to CE logo PNG (download from Case Engine Branding folder; recommend 350x180 dark variant)")
    parser.add_argument("--practice-area", required=True,
                        help="Practice area name (Title Case, e.g. 'Car Accidents')")
    parser.add_argument("--episode-topic", required=True,
                        help="Episode title (e.g. 'How to File a Car Accident Claim')")
    parser.add_argument("--scope", required=True,
                        choices=["Topic Only", "Location", "Extension"],
                        help="Scope of this template")
    parser.add_argument("--location", default=None,
                        help="State-prefixed jurisdictional folder name (e.g. 'GA - Savannah'); required for Location/Extension")
    parser.add_argument("--output", required=True,
                        help='Output DOCX path. Filename MUST follow the canonical pattern "ROS Template - {Episode Short Title} - {Location}.docx" so aggregated views (Drive search, recent files) stay scannable across templates. Example: <scope-folder>/ROS Template - How to File a Car Accident Claim - GA Savannah.docx. A sibling .md is written next to it.')
    parser.add_argument("--run-date", default=None,
                        help="Run date in 'Month D, YYYY' format (defaults to today)")
    args = parser.parse_args()

    if args.scope in ("Location", "Extension") and not args.location:
        print("ERROR: --location is required when --scope is Location or Extension", file=sys.stderr)
        sys.exit(1)

    build(args)


if __name__ == "__main__":
    main()
