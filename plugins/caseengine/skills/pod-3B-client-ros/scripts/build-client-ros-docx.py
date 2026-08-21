#!/usr/bin/env python3
"""
What: Builds a CE-branded Client ROS DOCX (and paired .md) from a client-ros-data.json.
Input: client-ros-data.json (filename per the canonical Podcast Drive doc:
       https://docs.google.com/document/d/1YhybGpp9DIqmV56P6OOHIQe7A6RxvjQGHLHKcUM0JmU)
       + a CE logo image (download from the canonical Case Engine Branding folder:
       https://drive.google.com/drive/folders/1OulNcg6hZ6caFD0Xc-W50i7ougC_y0qo)
Output:
  - {output}.docx — CE-branded Client ROS with cover page, headers/footers, segments.
  - {output}.md   — Plain-markdown sibling at the same base path. Pandoc artifacts are
                    stripped from the .md body (e.g. `[Houston PD]{.underline}` -> `Houston PD`)
                    while the .docx renders the same text as a true underlined run.
Re-run: Safe — overwrites both outputs. Drive's auto-conversion preserves the same fileId
        when the calling skill uses files.update for re-uploads (per the Push to Drive doc).

What this script does NOT include:
  - Any "Internal Setup" / "Complete and delete this section before sharing" checklist.
    That block is a Client Guide artifact only — never appears in the
    Client ROS Drive deliverable.
  - The Appendix material (Formatting Guide, extended Producer Notes, Entity Architecture,
    Search Queries) — those live in the upstream ROS Template, not in the populated
    Client ROS that ships to the firm's Drive folder.

Pandoc artifact handling (CRITICAL — Client ROS uses these heavily):
  - `[text]{.underline}` -> DOCX: real underlined run | MD: just `text`
  - `[text]{.smallcaps}` -> DOCX: small-caps run     | MD: just `text`
  - `[text]{.mark}`      -> DOCX: highlighted run    | MD: just `text`
  - `[text]{.color=...}` -> DOCX: colored run        | MD: just `text`
  Both **bold** and *italic* markdown markers ARE preserved in both outputs.

Convention sync: section order, branding spec, and the canonical filename live in the
Podcast Drive doc + the Case Engine Branding folder. If those change, update this script
to match. The canonical sources are the single source of truth.

Dependencies: python-docx (cowork runtime has it). No optional deps.

Usage:
  python3 build-client-ros-docx.py \\
      --data /path/to/client-ros-data.json \\
      --logo /path/to/ce-logo.png \\
      --firm "The May Firm" \\
      --attorney "Robert May" \\
      --show-name "Car Accident Attorney w. Robert May" \\
      --practice-area "Car Accidents" \\
      --scope "Location" \\
      --location "Santa Maria, California" \\
      --output "/path/to/Client ROS - E2 - How to File a Car Accident Claim - CA Santa Maria.docx" \\
      [--run-date "April 27, 2026"]   (defaults to today)
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# CE brand (per Case Engine Branding folder colors.md)
CE_BLUE = RGBColor(0x35, 0x73, 0xFF)
CE_DARK = RGBColor(0x0F, 0x17, 0x2A)
CE_GRAY = RGBColor(0x5B, 0x66, 0x76)
FONT = "Roboto"  # CE branded deliverable font (overrides Calibri per Gabe 2026-05-12; if the Branding folder spec still says Calibri, Roboto wins)


# ---------------------------------------------------------------------------
# Pandoc artifact regex set
# ---------------------------------------------------------------------------
# `[text]{.underline}` / `{.smallcaps}` / `{.mark}` / `{.color=...}`
PANDOC_SPAN_RE = re.compile(
    r"\[(?P<text>[^\]]+)\]\{(?P<attrs>[^}]+)\}"
)
# Legacy HTML <u>..</u> tags surfaced in older templates — treat the same way.
LEGACY_U_RE = re.compile(r"<u>(?P<text>.*?)</u>", re.DOTALL)


def parse_pandoc_attrs(attrs: str) -> dict:
    """Parse `.underline .smallcaps color=red` style attribute strings into flags."""
    flags = {"underline": False, "smallcaps": False, "mark": False, "color": None}
    for token in attrs.split():
        token = token.strip()
        if token == ".underline":
            flags["underline"] = True
        elif token == ".smallcaps":
            flags["smallcaps"] = True
        elif token == ".mark":
            flags["mark"] = True
        elif token.startswith(".color=") or token.startswith("color="):
            flags["color"] = token.split("=", 1)[1].strip('"').strip("'")
    return flags


def strip_pandoc_for_md(text: str) -> str:
    """Remove pandoc span wrappers so the .md body reads cleanly.

    `[Houston Police Department]{.underline}` -> `Houston Police Department`
    `<u>Houston Police Department</u>`        -> `Houston Police Department`
    Bold + italic markers survive untouched.
    """
    text = PANDOC_SPAN_RE.sub(lambda m: m.group("text"), text)
    text = LEGACY_U_RE.sub(lambda m: m.group("text"), text)
    return text


# ---------------------------------------------------------------------------
# Inline run rendering for the DOCX side
# ---------------------------------------------------------------------------
# Tokenize a string into segments of (text, bold, italic, underline, smallcaps, mark, color)
# so a single attorney bullet like
#   "**Safety first:** move toward the [shoulder]{.underline} and call *911*"
# renders with the right per-run styling in DOCX.

# Token-aware splitter for **bold**, *italic*, [text]{...}, <u>text</u>. Bold is
# non-greedy and can contain anything (including a nested pandoc span); we recurse into
# bold/italic inner content so `**[Firm Name]{.underline}**` renders as ONE run that is
# both bold AND underlined, never leaking the literal `**` markers or the
# `[...]{.underline}` text. Pandoc spans nested inside *italic* are handled the same way.
INLINE_TOKEN_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*)"                          # **bold** (non-greedy, may nest)
    r"|(\*(?P<ital>[^*]+?)\*)"                          # *italic*
    r"|(\[(?P<pinner>[^\]]+)\]\{(?P<pattrs>[^}]+)\})"   # [text]{.underline}
    r"|(<u>(?P<uinner>.*?)</u>)"                        # <u>legacy</u>
)


def tokenize_inline(text: str, *, bold=False, italic=False, underline=False,
                    smallcaps=False, mark=False, color=None):
    """Yield (text, bold, italic, underline, smallcaps, mark, color) tuples.

    Recursively descends into **bold** / *italic* spans so nested pandoc markers render
    as a single run carrying every applicable style. No literal markdown/pandoc/HTML
    markers leak to the rendered DOCX.
    """
    if not text:
        return
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        start, end = match.span()
        if start > pos:
            yield (text[pos:start], bold, italic, underline, smallcaps, mark, color)
        if match.group("bold") is not None:
            yield from tokenize_inline(match.group("bold"), bold=True, italic=italic,
                                       underline=underline, smallcaps=smallcaps,
                                       mark=mark, color=color)
        elif match.group("ital") is not None:
            yield from tokenize_inline(match.group("ital"), bold=bold, italic=True,
                                       underline=underline, smallcaps=smallcaps,
                                       mark=mark, color=color)
        elif match.group("pinner") is not None:
            attrs = parse_pandoc_attrs(match.group("pattrs"))
            yield from tokenize_inline(match.group("pinner"), bold=bold, italic=italic,
                                       underline=underline or attrs["underline"],
                                       smallcaps=smallcaps or attrs["smallcaps"],
                                       mark=mark or attrs["mark"],
                                       color=color or attrs["color"])
        elif match.group("uinner") is not None:
            yield from tokenize_inline(match.group("uinner"), bold=bold, italic=italic,
                                       underline=True, smallcaps=smallcaps,
                                       mark=mark, color=color)
        pos = end
    if pos < len(text):
        yield (text[pos:], bold, italic, underline, smallcaps, mark, color)


def set_run_style(run, *, font=FONT, size=11, bold=False, italic=False,
                  underline=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)


def add_inline_paragraph(doc, text, *, align=None, space_before=0, space_after=4,
                         base_size=11, base_color=CE_DARK, base_italic=False,
                         base_bold=False, list_bullet=False):
    """Render a paragraph that contains inline tokens (bold / italic / underline)."""
    if list_bullet:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for chunk, b, i, u, sc, mk, col in tokenize_inline(text):
        if not chunk:
            continue
        run = p.add_run(chunk)
        set_run_style(
            run, size=base_size,
            bold=base_bold or b,
            italic=base_italic or i,
            underline=u,
            color=base_color,
        )
        if sc:
            rPr = run._element.get_or_add_rPr()
            sc_el = OxmlElement("w:smallCaps")
            sc_el.set(qn("w:val"), "1")
            rPr.append(sc_el)
        if mk:
            rPr = run._element.get_or_add_rPr()
            hl = OxmlElement("w:highlight")
            hl.set(qn("w:val"), "yellow")
            rPr.append(hl)
    return p


def add_para(doc, text="", *, align=None, space_before=0, space_after=6, **run_kw):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_style(run, **run_kw)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_header(section, text):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_style(run, size=9, italic=True, color=CE_GRAY)


def set_footer(section, text_left):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text_left + "\t\t")
    set_run_style(run, size=9, color=CE_GRAY)
    run2 = p.add_run("Page ")
    set_run_style(run2, size=9, color=CE_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_style(run, size=16, bold=True, color=CE_BLUE)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_style(run, size=13, bold=True, color=CE_DARK)
    return p


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------
def build(args):
    data_path = Path(args.data)
    if not data_path.exists():
        print(f"ERROR: {data_path} not found", file=sys.stderr)
        sys.exit(1)
    data = json.loads(data_path.read_text())

    firm = args.firm
    attorney = args.attorney
    show_name = args.show_name
    practice_area = args.practice_area
    scope = args.scope
    location = args.location
    run_date = args.run_date or date.today().strftime("%B %-d, %Y")

    episode_topic = data.get("episode_topic", "")
    episode_number = data.get("episode_number", "")
    duration_min = data.get("duration_min", "")
    recording_date = data.get("recording_date", "")
    phone_number = data.get("phone_number", "")
    website = data.get("website", "")
    template_version = data.get("template_version", "")
    episode_goal = data.get("episode_goal", "authority")
    producer_notes = data.get("producer_notes", {})
    introduction = data.get("introduction", {})
    pre_show_checks = data.get("pre_show_checks", []) or []
    segments = data.get("segments", [])
    closing = data.get("closing", {})
    post_show_wrapup = data.get("post_show_wrapup", []) or []
    entity_checklist = data.get("entity_checklist", [])

    header_text = f"Case Engine  |  Client Run of Show  |  {firm}, {location}"

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        set_header(section, header_text)
        set_footer(section, "Case Engine  |  Confidential")

    # ---- Cover page ----
    for _ in range(3):
        doc.add_paragraph()

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(args.logo, width=Inches(2.4))

    add_para(doc, "Client Run of Show", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=6, size=36, bold=True, color=CE_BLUE)
    add_para(doc, episode_topic, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4, size=22, bold=True, color=CE_DARK)
    subtitle = f"{firm} - {location}" if location else firm
    add_para(doc, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=24, size=14, color=CE_DARK)
    add_para(doc, "Prepared by Case Engine", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=2, size=11, italic=True, color=CE_GRAY)
    add_para(doc, run_date, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=0, size=11, color=CE_GRAY)

    add_page_break(doc)

    # ---- Run of Show header ----
    add_h2(doc, f"Run of Show: {episode_topic}")
    add_inline_paragraph(
        doc,
        f"**Practice Area:** **{practice_area}**",
        space_after=2,
    )
    _ep_line = f"**Episode:** **{episode_number}**  |  **Duration:** ~{duration_min} minutes"
    if recording_date:
        _ep_line += f"  |  **Recording Date:** **{recording_date}**"
    add_inline_paragraph(
        doc,
        _ep_line,
        space_after=2,
    )
    add_inline_paragraph(
        doc,
        f"**Attorney:** **{attorney}**  |  **{firm}**  |  {location}  |  **{website}**",
        space_after=2,
    )
    add_inline_paragraph(
        doc,
        f"**Template Version:** {template_version}  |  **Episode Goal:** {episode_goal}",
        space_after=10,
    )

    # ---- Pre-Show Checks (optional) ----
    if pre_show_checks:
        add_h3(doc, "Pre-Show Checks")
        for line in pre_show_checks:
            add_inline_paragraph(doc, line, list_bullet=True)

    # ---- Producer Notes ----
    add_h2(doc, "Producer Notes")
    if producer_notes.get("jurisdiction"):
        add_inline_paragraph(
            doc, f"**Jurisdiction:** {producer_notes['jurisdiction']}",
            space_after=8,
        )
    add_inline_paragraph(
        doc, f"**Attorney website:** **{website}**",
        space_after=2,
    )
    if producer_notes.get("about_attorney"):
        add_inline_paragraph(
            doc, f"**About the attorney:** {producer_notes['about_attorney']}",
            space_after=8,
        )
    for flag in producer_notes.get("production_flags", []) or []:
        add_inline_paragraph(doc, f"**Production flag:** {flag}",
                             list_bullet=True)

    # ---- Introduction ----
    intro_dur = introduction.get("duration_min", 2)
    add_h2(doc, f"Introduction (~{intro_dur} minutes)")
    add_para(doc, "Interviewer", italic=True, color=CE_DARK, space_after=2)
    for para in introduction.get("paragraphs", []):
        add_inline_paragraph(doc, para, space_after=8)
    transition = introduction.get("transition_line") or "*Transition directly into Q1.*"
    add_inline_paragraph(doc, transition, space_after=8, base_italic=True)
    if introduction.get("co_host_notes"):
        add_inline_paragraph(
            doc, f"*Interviewer Notes: {introduction['co_host_notes']}*",
            space_after=8, base_italic=True, base_color=CE_GRAY,
        )

    # ---- Segments ----
    for segment in segments:
        seg_id = segment.get("segment_id", "")
        seg_name = segment.get("name", "")
        seg_dur = segment.get("duration_min", "")
        add_h2(doc, f"{seg_id}: {seg_name} (~{seg_dur} minutes)")
        if segment.get("intro_prompt"):
            add_inline_paragraph(doc, segment["intro_prompt"], space_after=8,
                                 base_italic=True, base_color=CE_GRAY)
        for q in segment.get("questions", []):
            q_id = q.get("q_id", "")
            q_text = q.get("q_text", "")
            q_dur = q.get("duration_min", "")
            add_h3(doc, f"{q_id}: {q_text} ({q_dur} min)")
            add_para(doc, "Interviewer", italic=True, color=CE_DARK, space_after=2)
            if q.get("co_host_setup"):
                add_inline_paragraph(doc, q["co_host_setup"], space_after=6)
            add_inline_paragraph(doc, f"**{q_text}**", space_after=6)
            add_para(doc, "Attorney", italic=True, color=CE_DARK,
                     space_after=2)
            for bullet in q.get("attorney_bullets", []):
                add_inline_paragraph(doc, bullet, list_bullet=True)
        if segment.get("segment_wrap"):
            add_para(doc, "Interviewer", italic=True, color=CE_DARK, space_after=2)
            add_inline_paragraph(doc, segment["segment_wrap"], space_after=10)

    # ---- Closing and Call to Action ----
    closing_dur = closing.get("duration_min", 2)
    add_h2(doc, f"Closing and Call to Action (~{closing_dur} minutes)")
    add_para(doc, "Interviewer", italic=True, color=CE_DARK, space_after=2)
    if closing.get("final_takeaway_question"):
        # The whole question is bolded by the wrap below; strip any inner/partial
        # bold markers first so a field that already carries `**...**` does not
        # double-wrap and leak literal `**` (Spaulding Ep6, 2026-07-06).
        _ftq = closing["final_takeaway_question"].replace("**", "")
        add_inline_paragraph(doc, f"**{_ftq}**", space_after=6)
    add_para(doc, "Attorney", italic=True, color=CE_DARK,
             space_after=2)
    for bullet in closing.get("attorney_takeaway_bullets", []):
        add_inline_paragraph(doc, bullet, list_bullet=True)
    if closing.get("cta_line"):
        add_inline_paragraph(doc, closing["cta_line"], space_after=10)

    # ---- Post-Show Wrap-up (optional) ----
    if post_show_wrapup:
        add_h2(doc, "Post-Show Wrap-up")
        for line in post_show_wrapup:
            add_inline_paragraph(doc, line, list_bullet=True)

    add_para(doc, "End of Run of Show", align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=12, space_after=12, size=11, italic=True, color=CE_GRAY)

    # ---- Entity Checklist ----
    if entity_checklist:
        add_h2(doc, "Entity Checklist")
        add_para(doc,
                 "Recording-time tally sheet. Producer fills the Actual Mentions column during post-record review.",
                 size=11, italic=True, color=CE_GRAY, space_after=8)
        from docx.enum.table import WD_TABLE_ALIGNMENT
        headers = ["Entity", "Questions", "Target Mentions", "Actual Mentions"]
        table = doc.add_table(rows=1 + len(entity_checklist), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(h)
            set_run_style(run, size=10, bold=True, color=CE_BLUE)
        for r_idx, row in enumerate(entity_checklist):
            cells = table.rows[1 + r_idx].cells
            cells[0].text = ""
            # Render entity with inline tokens (preserves underline)
            p = cells[0].paragraphs[0]
            for chunk, b, i, u, sc, mk, col in tokenize_inline(row.get("entity", "")):
                if not chunk:
                    continue
                run = p.add_run(chunk)
                set_run_style(run, size=10, bold=b, italic=i, underline=u,
                              color=CE_DARK)
            questions = row.get("questions", [])
            if isinstance(questions, list):
                questions = ", ".join(questions)
            # Tokenize the remaining cells too (matches the entity cell) so any
            # stray markdown (**bold**, [x]{.underline}) renders as real runs and
            # never leaks literal markers into the table.
            for ci, val in ((1, str(questions)),
                            (2, str(row.get("target_mentions", ""))),
                            (3, str(row.get("actual_mentions", "")))):
                cells[ci].text = ""
                p = cells[ci].paragraphs[0]
                for chunk, b, i, u, sc, mk, col in tokenize_inline(val):
                    if not chunk:
                        continue
                    run = p.add_run(chunk)
                    set_run_style(run, size=10, bold=b, italic=i, underline=u,
                                  color=CE_DARK)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved DOCX: {output_path}")

    # ---- Markdown sibling ----
    md_path = output_path.with_suffix(".md")
    md_text = build_markdown(
        firm=firm, attorney=attorney, show_name=show_name,
        practice_area=practice_area, scope=scope, location=location,
        run_date=run_date, episode_topic=episode_topic,
        episode_number=episode_number, duration_min=duration_min,
        recording_date=recording_date, phone_number=phone_number,
        website=website, template_version=template_version,
        episode_goal=episode_goal, producer_notes=producer_notes,
        introduction=introduction, pre_show_checks=pre_show_checks,
        segments=segments, closing=closing, post_show_wrapup=post_show_wrapup,
        entity_checklist=entity_checklist,
    )
    md_path.write_text(md_text)
    print(f"Saved MD:   {md_path}")
    seg_count = len(segments)
    q_count = sum(len(s.get("questions", [])) for s in segments)
    print(f"Data: {seg_count} segments, {q_count} questions, "
          f"{len(entity_checklist)} entity-checklist rows")


def build_markdown(*, firm, attorney, show_name, practice_area, scope, location,
                   run_date, episode_topic, episode_number, duration_min,
                   recording_date, phone_number, website, template_version,
                   episode_goal, producer_notes, introduction, pre_show_checks,
                   segments, closing, post_show_wrapup, entity_checklist):
    """Build the raw .md deliverable.

    Pandoc artifacts ([text]{.underline}, etc.) are stripped to plain text in
    the .md body — bold and italic markdown markers ARE preserved. The .docx
    sibling carries the same content with proper formatting runs.
    """
    lines = []

    lines.append(f"# Run of Show: {episode_topic}")
    lines.append("")
    lines.append(f"**Practice Area:** **{practice_area}**")
    lines.append(
        f"**Episode:** **{episode_number}** | **Duration:** ~{duration_min} minutes"
    )
    if recording_date:
        lines.append(f"**Recording Date:** **{recording_date}**")
    if template_version:
        lines.append(f"**Template Version:** {template_version}")
    lines.append(f"**Location:** {location}" if location else "")
    lines.append(f"**Episode Goal:** {episode_goal}")
    lines.append("")
    lines.append(f"_Prepared by Case Engine - {run_date}_")
    lines.append("")
    lines.append("---")
    lines.append("")

    if pre_show_checks:
        lines.append("## Pre-Show Checks")
        lines.append("")
        for line in pre_show_checks:
            lines.append(f"- {strip_pandoc_for_md(line)}")
        lines.append("")

    lines.append("## Producer Notes")
    lines.append("")
    if producer_notes.get("jurisdiction"):
        lines.append(
            f"**Jurisdiction:** {strip_pandoc_for_md(producer_notes['jurisdiction'])}"
        )
        lines.append("")
    lines.append(f"**Attorney website:** **{website}**")
    if producer_notes.get("about_attorney"):
        lines.append(
            f"**About the attorney:** {strip_pandoc_for_md(producer_notes['about_attorney'])}"
        )
    for flag in producer_notes.get("production_flags", []) or []:
        lines.append(f"- **Production flag:** {strip_pandoc_for_md(flag)}")
    lines.append("")
    lines.append("---")
    lines.append("")

    intro_dur = introduction.get("duration_min", 2)
    lines.append(f"## Introduction (~{intro_dur} minutes)")
    lines.append("")
    lines.append("*Interviewer*")
    lines.append("")
    for para in introduction.get("paragraphs", []):
        lines.append(strip_pandoc_for_md(para))
        lines.append("")
    transition = introduction.get("transition_line") or "*Transition directly into Q1.*"
    lines.append(strip_pandoc_for_md(transition))
    lines.append("")
    if introduction.get("co_host_notes"):
        lines.append(
            f"*Interviewer Notes: {strip_pandoc_for_md(introduction['co_host_notes'])}*"
        )
        lines.append("")
    lines.append("---")
    lines.append("")

    for segment in segments:
        seg_id = segment.get("segment_id", "")
        seg_name = segment.get("name", "")
        seg_dur = segment.get("duration_min", "")
        lines.append(f"## {seg_id}: {seg_name} (~{seg_dur} minutes)")
        lines.append("")
        if segment.get("intro_prompt"):
            lines.append(f"_{strip_pandoc_for_md(segment['intro_prompt'])}_")
            lines.append("")
        for q in segment.get("questions", []):
            q_id = q.get("q_id", "")
            q_text = q.get("q_text", "")
            q_dur = q.get("duration_min", "")
            lines.append(f"### {q_id}: {q_text} ({q_dur} min)")
            lines.append("")
            lines.append("*Interviewer*")
            lines.append("")
            if q.get("co_host_setup"):
                lines.append(strip_pandoc_for_md(q["co_host_setup"]))
                lines.append("")
            lines.append(f"**{strip_pandoc_for_md(q_text)}**")
            lines.append("")
            lines.append("*Attorney*")
            lines.append("")
            for bullet in q.get("attorney_bullets", []):
                lines.append(f"- {strip_pandoc_for_md(bullet)}")
            lines.append("")
        if segment.get("segment_wrap"):
            lines.append("*Interviewer*")
            lines.append("")
            lines.append(strip_pandoc_for_md(segment["segment_wrap"]))
            lines.append("")
        lines.append("---")
        lines.append("")

    closing_dur = closing.get("duration_min", 2)
    lines.append(f"## Closing and Call to Action (~{closing_dur} minutes)")
    lines.append("")
    lines.append("*Interviewer*")
    lines.append("")
    if closing.get("final_takeaway_question"):
        _ftq = strip_pandoc_for_md(closing["final_takeaway_question"]).replace("**", "")
        lines.append(f"**{_ftq}**")
        lines.append("")
    lines.append("*Attorney*")
    lines.append("")
    for bullet in closing.get("attorney_takeaway_bullets", []):
        lines.append(f"- {strip_pandoc_for_md(bullet)}")
    lines.append("")
    if closing.get("cta_line"):
        lines.append(strip_pandoc_for_md(closing["cta_line"]))
        lines.append("")

    if post_show_wrapup:
        lines.append("## Post-Show Wrap-up")
        lines.append("")
        for line in post_show_wrapup:
            lines.append(f"- {strip_pandoc_for_md(line)}")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("*End of Run of Show*")
    lines.append("")

    if entity_checklist:
        lines.append("## Entity Checklist")
        lines.append("")
        lines.append("| Entity | Questions | Target Mentions | Actual Mentions |")
        lines.append("|---|---|---|---|")
        for row in entity_checklist:
            entity = strip_pandoc_for_md(row.get("entity", ""))
            questions = row.get("questions", [])
            if isinstance(questions, list):
                questions = ", ".join(questions)
            target = row.get("target_mentions", "")
            actual = row.get("actual_mentions", "")
            lines.append(f"| {entity} | {questions} | {target} | {actual} |")
        lines.append("")

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Build CE-branded Client ROS DOCX (and paired .md) from client-ros-data.json"
    )
    parser.add_argument("--data", required=True,
                        help="Path to client-ros-data.json")
    parser.add_argument("--logo", required=True,
                        help="Path to CE logo PNG (download from Case Engine Branding folder; 350x180 dark variant recommended)")
    parser.add_argument("--firm", required=True, help="Firm name (e.g., 'The May Firm')")
    parser.add_argument("--attorney", required=True,
                        help="Attorney full name (e.g., 'Robert May')")
    parser.add_argument("--show-name", required=True,
                        help="Podcast show name (e.g., 'Car Accident Attorney w. Robert May')")
    parser.add_argument("--practice-area", required=True,
                        help="Practice area (Title Case, e.g. 'Car Accidents')")
    parser.add_argument("--scope", required=True,
                        choices=["Topic Only", "Location", "Extension"],
                        help="Scope of this Client ROS")
    parser.add_argument("--location", default="",
                        help="Location string for cover / header (e.g., 'Santa Maria, California'). Empty for Topic Only.")
    parser.add_argument("--output", required=True,
                        help='Output DOCX path - write it into the firm episode folder\'s "01 Strategy/" directory ALONGSIDE the ROS Template (already there, read-only) and the Client Guide (lands later, RoS Step 4). Filename pattern "Client ROS - E{N} - {Episode Short Title} - {Location}.docx" (append " (Extension)" for an extension cell). Example: "<episode-folder>/01 Strategy/Client ROS - E2 - How to File a Car Accident Claim - CA Santa Maria.docx". A sibling .md is written next to the DOCX. The DOCX is uploaded to Drive as a Google-Doc mimeType so Drive auto-converts to a clean branded Doc with real underlined entity runs - never re-upload the .md with convert=true (that leaks "[entity]{.underline}" as visible text).')
    parser.add_argument("--run-date", default=None,
                        help="Run date in 'Month D, YYYY' format (defaults to today)")
    args = parser.parse_args()

    if args.scope in ("Location", "Extension") and not args.location:
        print("ERROR: --location is required when --scope is Location or Extension",
              file=sys.stderr)
        sys.exit(1)

    build(args)


if __name__ == "__main__":
    main()
